#!/usr/bin/env python3
"""
FastTracker Algorithm Description Document Generator

Generates a comprehensive Japanese algorithm description document
(algorithm_description.docx) for the FastTracker system.

Usage:
    python generate_algorithm_doc.py

Requirements:
    - python-docx
    - matplotlib
"""

import os
import sys
import tempfile
import shutil

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.mathtext

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DOCX = os.path.join(SCRIPT_DIR, "algorithm_description.docx")
TEMP_DIR = tempfile.mkdtemp(prefix="fasttracker_doc_")

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
FONT_NAME_JP = "游ゴシック"
FONT_NAME_EN = "Yu Gothic"
TITLE_COLOR = RGBColor(0x1A, 0x3C, 0x6E)  # Dark navy
HEADING_COLOR = RGBColor(0x1F, 0x4E, 0x79)
ACCENT_COLOR = RGBColor(0x2E, 0x75, 0xB6)
TABLE_HEADER_BG = "1F4E79"
TABLE_ALT_BG = "D6E4F0"

# Track generated images for cleanup
_image_counter = 0


def render_math(latex_str, filename=None, fontsize=14, dpi=200):
    """Render a LaTeX math string to a PNG image using matplotlib mathtext.

    Args:
        latex_str: LaTeX math string (without surrounding $).
        filename: Output filename (auto-generated if None).
        fontsize: Font size for the rendered equation.
        dpi: Resolution of the output image.

    Returns:
        Absolute path to the generated PNG file.
    """
    global _image_counter
    if filename is None:
        _image_counter += 1
        filename = f"eq_{_image_counter:03d}.png"

    filepath = os.path.join(TEMP_DIR, filename)

    fig = plt.figure(figsize=(0.01, 0.01))
    fig.patch.set_alpha(0.0)

    text = fig.text(
        0, 0,
        f"${latex_str}$",
        fontsize=fontsize,
        color="black",
        ha="left",
        va="baseline",
    )

    fig.savefig(
        filepath,
        dpi=dpi,
        transparent=True,
        bbox_inches="tight",
        pad_inches=0.05,
    )
    plt.close(fig)
    return filepath


def render_math_block(latex_str, filename=None, fontsize=16, dpi=200):
    """Render a display-style math block (larger, centered)."""
    return render_math(latex_str, filename=filename, fontsize=fontsize, dpi=dpi)


def save_figure(fig, filename=None):
    """Save a matplotlib figure to the temp directory and close it.

    Args:
        fig: matplotlib Figure object.
        filename: Output filename (auto-generated if None).

    Returns:
        Absolute path to the generated PNG file.
    """
    global _image_counter
    if filename is None:
        _image_counter += 1
        filename = f"fig_{_image_counter:03d}.png"
    filepath = os.path.join(TEMP_DIR, filename)
    fig.savefig(filepath, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return filepath


def add_figure_to_doc(doc, fig_path, caption=None, width_inches=5.5):
    """Insert a saved figure PNG into the document with an optional caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(fig_path, width=Inches(width_inches))
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(8)
        run = cap_p.add_run(caption)
        set_run_font(run, size=9, italic=True, color=RGBColor(0x40, 0x40, 0x40))


def add_table_caption(doc, caption):
    """Add a bold table caption paragraph (above the table)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(caption)
    set_run_font(run, size=9, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))


def add_symbol_legend(doc, entries):
    """Add a compact 4-column symbol/definition legend table after equations.

    Args:
        doc:     Document object.
        entries: list of (symbol_str, definition_str) tuples.
    """
    # Section label
    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_before = Pt(2)
    p_hdr.paragraph_format.space_after = Pt(1)
    run_hdr = p_hdr.add_run("【記号の定義】")
    set_run_font(run_hdr, size=8.5, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

    n = len(entries)
    half = (n + 1) // 2  # number of rows

    tbl = doc.add_table(rows=half, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i in range(half):
        for pair_offset, entry_idx in [(0, i), (2, i + half)]:
            if entry_idx < n:
                sym, defn = entries[entry_idx]
                sym_cell = tbl.rows[i].cells[pair_offset]
                def_cell = tbl.rows[i].cells[pair_offset + 1]

                set_cell_shading(sym_cell, "D6E8FF")
                p_s = sym_cell.paragraphs[0]
                p_s.paragraph_format.space_before = Pt(1)
                p_s.paragraph_format.space_after = Pt(1)
                run_s = p_s.add_run(sym)
                set_run_font(run_s, size=8.5, bold=True,
                             color=RGBColor(0x1F, 0x4E, 0x79))

                set_cell_shading(def_cell, "F4F8FF")
                p_d = def_cell.paragraphs[0]
                p_d.paragraph_format.space_before = Pt(1)
                p_d.paragraph_format.space_after = Pt(1)
                run_d = p_d.add_run(defn)
                set_run_font(run_d, size=8.5)

    # Column widths: sym=1.8cm, def=4.8cm (×2 pairs)
    for row in tbl.rows:
        row.cells[0].width = Cm(1.8)
        row.cells[1].width = Cm(4.8)
        row.cells[2].width = Cm(1.8)
        row.cells[3].width = Cm(4.8)

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, size=10, bold=False, italic=False, color=None, name=None):
    """Apply font formatting to a run."""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    fname = name or FONT_NAME_JP
    run.font.name = fname
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), fname)


def add_paragraph_with_font(doc, text, style=None, size=10, bold=False,
                            alignment=None, space_after=Pt(6), color=None):
    """Add a paragraph with proper Japanese font settings."""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after

    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_math_image(doc, latex_str, fontsize=14, width_inches=None):
    """Render a math formula and insert it as an inline image."""
    img_path = render_math(latex_str, fontsize=fontsize)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    if width_inches:
        run.add_picture(img_path, width=Inches(width_inches))
    else:
        # Auto-size but cap max width
        from PIL import Image as PILImage
        try:
            with PILImage.open(img_path) as im:
                w_px, h_px = im.size
            # Limit to 6 inches wide
            w_in = min(w_px / 200.0, 6.0)
            run.add_picture(img_path, width=Inches(w_in))
        except Exception:
            run.add_picture(img_path, width=Inches(4.0))
    return p


def add_math_inline(paragraph, latex_str, fontsize=12):
    """Add an inline math image to an existing paragraph."""
    img_path = render_math(latex_str, fontsize=fontsize)
    run = paragraph.add_run()
    try:
        from PIL import Image as PILImage
        with PILImage.open(img_path) as im:
            w_px, h_px = im.size
        w_in = min(w_px / 200.0, 3.0)
        run.add_picture(img_path, width=Inches(w_in))
    except Exception:
        run.add_picture(img_path, width=Inches(2.0))


def add_bullet(doc, text, level=0, size=10):
    """Add a bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p


def add_table_with_header(doc, headers, rows, col_widths=None):
    """Add a formatted table with a colored header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        set_cell_shading(cell, TABLE_HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        set_run_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[1 + row_idx]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            if row_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_run_font(run, size=9)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table


def setup_styles(doc):
    """Configure document styles for Japanese text."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME_JP
    font.size = Pt(10)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME_JP)

    # Heading styles
    for level in range(1, 4):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_font = h_style.font
            h_font.name = FONT_NAME_JP
            h_font.color.rgb = HEADING_COLOR
            h_rPr = h_style.element.get_or_add_rPr()
            h_rFonts = h_rPr.find(qn("w:rFonts"))
            if h_rFonts is None:
                h_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                h_rPr.insert(0, h_rFonts)
            h_rFonts.set(qn("w:eastAsia"), FONT_NAME_JP)

            if level == 1:
                h_font.size = Pt(18)
            elif level == 2:
                h_font.size = Pt(14)
            else:
                h_font.size = Pt(12)


def setup_page(doc):
    """Set up A4 page with reasonable margins and page numbers."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    # Add page numbers in footer
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Page number field
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar2)

    run4 = p.add_run("1")
    set_run_font(run4, size=9, color=RGBColor(0x80, 0x80, 0x80))

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar3)


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def add_title_page(doc):
    """Create a professional title page."""
    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FastTracker")
    set_run_font(run, size=36, bold=True, color=TITLE_COLOR, name="Consolas")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("アルゴリズム解説書")
    set_run_font(run2, size=28, bold=True, color=TITLE_COLOR)

    # Subtitle
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("GPU加速型マルチターゲット追尾システム")
    set_run_font(run3, size=14, color=ACCENT_COLOR)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("IMM-UKF / CUDA / 弾道ミサイル・HGV対応")
    set_run_font(run4, size=12, color=ACCENT_COLOR)

    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    # Document info table
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("文書番号", "FT-ALG-001"),
        ("版数", "1.0"),
        ("作成日", "2026年2月17日"),
        ("分類", "技術文書"),
    ]
    for i, (label, value) in enumerate(info_data):
        cell_l = info_table.rows[i].cells[0]
        cell_l.text = ""
        set_cell_shading(cell_l, TABLE_HEADER_BG)
        p = cell_l.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        cell_l.width = Cm(4)

        cell_r = info_table.rows[i].cells[1]
        cell_r.text = ""
        p = cell_r.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        set_run_font(run, size=10)
        cell_r.width = Cm(8)

    doc.add_page_break()


def add_toc(doc):
    """Add a table of contents page."""
    doc.add_heading("目次", level=1)

    # TOC field code
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        ' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar2)

    # Placeholder text
    run4 = p.add_run("（目次はWordで「フィールドの更新」を実行すると自動生成されます）")
    set_run_font(run4, size=10, italic=True, color=RGBColor(0x80, 0x80, 0x80))

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar3)

    # Manual TOC listing for reference
    doc.add_paragraph()
    toc_entries = [
        ("1.", "システム概要"),
        ("2.", "ターゲットジェネレータ"),
        ("  2.1", "軌道タイプ"),
        ("  2.2", "目標状態ベクトル"),
        ("  2.3", "弾道ミサイルモデル（RK4）"),
        ("  2.4", "HGV（極超音速滑空体）モデル"),
        ("3.", "レーダシミュレータ"),
        ("  3.1", "観測モデル"),
        ("  3.2", "Swerling II 検出モデル"),
        ("  3.3", "検出処理チェーン"),
        ("  3.4", "クラッタモデル"),
        ("4.", "状態空間モデル"),
        ("5.", "Unscented Kalman Filter (UKF)"),
        ("6.", "IMM (Interacting Multiple Model) フィルタ"),
        ("7.", "運動モデル"),
        ("  7.1", "CV（等速直線）モデル"),
        ("  7.2", "弾道（RK4）モデル"),
        ("  7.3", "CT（旋回）モデル"),
        ("8.", "データアソシエーション"),
        ("9.", "航跡管理"),
        ("10.", "評価指標"),
        ("11.", "GPU並列処理による高速化"),
        ("  11.1", "UKFシグマポイントの並列化"),
        ("  11.2", "IMMフィルタのストリーム並列化"),
        ("  11.3", "コスト行列計算の並列化"),
    ]
    for num, title in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{num}  {title}")
        indent = "  " in num
        set_run_font(run, size=10 if not indent else 9.5)

    doc.add_page_break()


def add_section_1(doc):
    """Section 1: System Overview."""
    doc.add_heading("1. システム概要", level=1)

    add_paragraph_with_font(
        doc,
        "FastTrackerは、弾道ミサイルおよび極超音速滑空体（HGV）を対象とした"
        "GPU加速型マルチターゲット追尾システムである。CUDAによる並列計算を活用し、"
        "UKFシグマポイント演算およびコスト行列計算を高速化する。",
        size=10,
    )

    doc.add_heading("1.1 システムアーキテクチャ", level=2)

    add_paragraph_with_font(
        doc,
        "本システムは6つの主要コンポーネントで構成される。"
        "各コンポーネントはパイプライン構造で逐次処理を行う（表1-1）。",
        size=10,
    )

    # Architecture pipeline
    pipeline = [
        ("Target Generator", "目標生成器", "弾道ミサイル・HGV等の目標軌道を物理モデルに基づき生成"),
        ("Radar Simulator", "レーダシミュレータ", "観測ノイズ・検出確率を模擬し、観測値を生成"),
        ("Data Association", "データアソシエーション", "観測値と既存航跡の対応付け（ハンガリアン法）"),
        ("IMM-UKF", "IMM-UKFフィルタ", "複数運動モデルによる状態推定（GPU加速）"),
        ("Track Manager", "航跡管理器", "航跡の生成・確認・削除を管理"),
        ("Evaluator", "評価器", "RMSE, OSPA等の追尾性能評価"),
    ]

    add_table_caption(doc, "表1-1  システムコンポーネント一覧")
    add_table_with_header(
        doc,
        ["コンポーネント", "日本語名", "機能概要"],
        [(c, j, d) for c, j, d in pipeline],
        col_widths=[4.0, 4.0, 8.0],
    )

    doc.add_paragraph()
    add_paragraph_with_font(
        doc,
        "処理フロー:",
        size=10, bold=True,
    )
    add_paragraph_with_font(
        doc,
        "Target Generator → Radar Simulator → Data Association → IMM-UKF → Track Manager → Evaluator",
        size=10,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.add_heading("1.2 GPU加速", level=2)
    add_paragraph_with_font(
        doc,
        "CUDAを用いた以下の処理のGPU並列化により、リアルタイム処理を実現する。",
        size=10,
    )
    add_bullet(doc, "UKFシグマポイントの生成・伝播（各航跡を並列処理）")
    add_bullet(doc, "コスト行列の計算（観測-航跡ペアを並列計算）")
    add_bullet(doc, "行列演算（コレスキー分解、行列乗算等）")

    doc.add_page_break()


def _make_trajectory_figure():
    """Create a figure illustrating three trajectory types."""
    import numpy as np
    fig, ax = plt.subplots(figsize=(7, 4))
    t = np.linspace(0, 1, 300)

    # Ballistic missile arc
    x_bal = t * 1200
    y_bal = 500 * np.sin(np.pi * t)
    ax.plot(x_bal, y_bal, color="#E05A2B", lw=2.5, label="Ballistic Missile")

    # HGV glide trajectory
    x_hgv = t * 900
    y_hgv = 80 * (1 - t * 0.7) + 20 * np.sin(3 * np.pi * t * 0.6)
    y_hgv = np.clip(y_hgv, 20, 120)
    ax.plot(x_hgv, y_hgv, color="#2E75B6", lw=2.5, linestyle="--",
            label="HGV (Hypersonic Glide)")

    # Constant velocity (cruise)
    x_cv = t * 600
    y_cv = np.ones_like(t) * 15
    ax.plot(x_cv, y_cv, color="#2EB67D", lw=2.0, linestyle=":",
            label="Constant Velocity")

    ax.set_xlabel("Horizontal Distance (km)", fontsize=10)
    ax.set_ylabel("Altitude (km)", fontsize=10)
    ax.set_title("Trajectory Types", fontsize=12, fontweight="bold")
    ax.legend(fontsize=9, loc="upper right")
    ax.set_xlim(0, 1300)
    ax.set_ylim(-10, 560)
    ax.grid(True, alpha=0.3)
    ax.set_facecolor("#F8F9FA")
    fig.tight_layout()
    return fig


def _make_ballistic_phases_figure():
    """Create a figure showing the three flight phases of a ballistic missile."""
    import numpy as np
    fig, ax = plt.subplots(figsize=(7, 3.5))
    t = np.linspace(0, 1, 300)
    x_bal = t * 1200
    y_bal = 500 * np.sin(np.pi * t)
    ax.plot(x_bal, y_bal, color="#E05A2B", lw=2.5)

    # Phase regions
    boost_end = 0.15
    terminal_start = 0.85
    ax.axvspan(0, boost_end * 1200, alpha=0.12, color="#FF8C00", label="Boost Phase")
    ax.axvspan(boost_end * 1200, terminal_start * 1200, alpha=0.10, color="#4472C4",
               label="Midcourse Phase (RK4)")
    ax.axvspan(terminal_start * 1200, 1200, alpha=0.12, color="#70AD47",
               label="Terminal Phase")

    # Labels
    ax.text(boost_end * 1200 * 0.5, 30, "Boost", ha="center", fontsize=8,
            color="#CC6600", fontweight="bold")
    ax.text((boost_end + terminal_start) * 600, 460, "Midcourse\n(Gravity+Drag)", ha="center",
            fontsize=8, color="#2E4D8C", fontweight="bold")
    ax.text((terminal_start + 1) * 600, 30, "Terminal", ha="center", fontsize=8,
            color="#4D7C1A", fontweight="bold")

    ax.set_xlabel("Horizontal Distance (km)", fontsize=10)
    ax.set_ylabel("Altitude (km)", fontsize=10)
    ax.set_title("Ballistic Missile Flight Phases", fontsize=12, fontweight="bold")
    ax.legend(fontsize=9, loc="upper right")
    ax.set_xlim(0, 1200)
    ax.set_ylim(-10, 530)
    ax.grid(True, alpha=0.3)
    ax.set_facecolor("#F8F9FA")
    fig.tight_layout()
    return fig


def _make_swerling_pd_figure():
    """P(D) vs SNR_avg curves for Swerling II model at multiple P_FA values."""
    import numpy as np
    fig, ax = plt.subplots(figsize=(6.5, 4))
    snr_dB = np.linspace(-2, 22, 400)
    snr_lin = 10 ** (snr_dB / 10.0)
    pfa_list = [1e-2, 1e-4, 1e-6, 1e-8]
    colors = ["#E05A2B", "#2E75B6", "#2EB67D", "#9B59B6"]
    for pfa, color in zip(pfa_list, colors):
        gamma_T = -np.log(pfa)
        pd = np.exp(-gamma_T / snr_lin)
        ax.plot(snr_dB, pd, lw=2.0, color=color,
                label=f"$P_{{FA}}=10^{{{int(round(np.log10(pfa)))}}}$")
    ax.set_xlabel("Mean SNR (dB)", fontsize=10)
    ax.set_ylabel("Detection Probability $P_D$", fontsize=10)
    ax.set_title("Swerling II: $P_D$ vs Mean SNR", fontsize=12, fontweight="bold")
    ax.set_xlim(-2, 22)
    ax.set_ylim(0, 1.02)
    ax.legend(fontsize=9)
    ax.grid(True, alpha=0.3)
    ax.axhline(0.5, color="gray", lw=0.8, linestyle=":")
    ax.axhline(0.9, color="gray", lw=0.8, linestyle=":")
    ax.set_facecolor("#F8F9FA")
    fig.tight_layout()
    return fig


def _make_detection_chain_figure():
    """Create a flowchart of the radar detection chain."""
    fig, ax = plt.subplots(figsize=(7, 3.2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3)
    ax.axis("off")
    ax.set_facecolor("white")

    boxes = [
        (0.3, "FOV\nCheck", "#4472C4"),
        (2.0, "Beam\nDirection", "#4472C4"),
        (3.7, "SNR_avg\nCalc.", "#2E75B6"),
        (5.4, "Exp. RV\nGen.", "#2EB67D"),
        (7.1, "CFAR\nThreshold", "#E05A2B"),
    ]

    from matplotlib.patches import FancyBboxPatch
    box_w, box_h = 1.4, 1.0
    cy = 1.5
    for (cx, label, color) in boxes:
        rect = FancyBboxPatch((cx, cy - box_h / 2), box_w, box_h,
                              boxstyle="round,pad=0.05",
                              facecolor=color, edgecolor="white",
                              linewidth=1.5, alpha=0.90)
        ax.add_patch(rect)
        ax.text(cx + box_w / 2, cy, label, ha="center", va="center",
                fontsize=7.5, color="white", fontweight="bold", linespacing=1.4)

    # Arrows
    arrow_props = dict(arrowstyle="-|>", color="#555555", lw=1.2)
    xs_end = [(cx + box_w) for (cx, _, _) in boxes[:-1]]
    xs_start = [cx for (cx, _, _) in boxes[1:]]
    for x0, x1 in zip(xs_end, xs_start):
        ax.annotate("", xy=(x1, cy), xytext=(x0, cy),
                    arrowprops=dict(arrowstyle="-|>", color="#555555", lw=1.5))

    # Result labels
    ax.text(9.0, cy + 0.65, "Detect", ha="center", fontsize=9,
            color="#2EB67D", fontweight="bold")
    ax.text(9.0, cy - 0.65, "Miss", ha="center", fontsize=9,
            color="#999999")
    ax.annotate("", xy=(9.2, cy + 0.4), xytext=(8.5 + box_w * 0.5 - 0.15, cy),
                arrowprops=dict(arrowstyle="-|>", color="#2EB67D", lw=1.4))
    ax.annotate("", xy=(9.2, cy - 0.4), xytext=(8.5 + box_w * 0.5 - 0.15, cy),
                arrowprops=dict(arrowstyle="-|>", color="#999999", lw=1.4))

    # SNR formula note
    ax.text(4.4, 0.35, r"$\mathrm{SNR}_{inst} = \mathrm{SNR}_{avg} \cdot \mathrm{Exp}(1)$",
            ha="center", fontsize=8, color="#444444",
            style="italic")
    ax.text(7.8, 0.35, r"$\gamma_T = -\ln(P_{FA})$",
            ha="center", fontsize=8, color="#444444",
            style="italic")

    fig.tight_layout()
    return fig


def _make_hgv_forces_figure():
    """Diagram showing HGV lift, drag, and bank angle in 3D concept."""
    import numpy as np
    fig, axes = plt.subplots(1, 2, figsize=(7.5, 3.2))

    # --- Left: side view with flight path angle γ ---
    ax = axes[0]
    ax.set_xlim(-0.5, 5.5)
    ax.set_ylim(-1.0, 3.5)
    ax.axis("off")
    ax.set_facecolor("#F8F9FA")
    ax.set_title("Side View (γ: flight path angle)", fontsize=9, fontweight="bold")

    from matplotlib.patches import FancyArrowPatch
    # Trajectory arc
    t_arc = np.linspace(0, 0.8, 50)
    x_arc = t_arc * 4.0
    y_arc = 1.8 + 0.4 * np.sin(np.pi * t_arc / 0.8)
    ax.plot(x_arc, y_arc, color="#E05A2B", lw=2.0, linestyle="--", label="Trajectory")

    # At midpoint: draw velocity vector, drag, lift, gravity
    mid = 25
    px, py = x_arc[mid], y_arc[mid]
    # Velocity direction (tangent to arc)
    dx_v = (x_arc[mid+1] - x_arc[mid-1])
    dy_v = (y_arc[mid+1] - y_arc[mid-1])
    spd = np.sqrt(dx_v**2 + dy_v**2)
    dx_v /= spd; dy_v /= spd

    # Velocity vector
    ax.annotate("", xy=(px + 0.9*dx_v, py + 0.9*dy_v), xytext=(px, py),
                arrowprops=dict(arrowstyle="-|>", color="#2E75B6", lw=2.0))
    ax.text(px + 0.9*dx_v + 0.05, py + 0.9*dy_v + 0.05, "v", color="#2E75B6",
            fontsize=9, fontweight="bold")

    # Drag vector (opposite to velocity)
    ax.annotate("", xy=(px - 0.7*dx_v, py - 0.7*dy_v), xytext=(px, py),
                arrowprops=dict(arrowstyle="-|>", color="#E05A2B", lw=2.0))
    ax.text(px - 0.85*dx_v - 0.15, py - 0.7*dy_v, r"$F_{drag}$", color="#E05A2B", fontsize=8)

    # Lift vector (perpendicular to velocity, upward)
    lx, ly = -dy_v, dx_v  # Perpendicular (upward component)
    ax.annotate("", xy=(px + 0.8*lx, py + 0.8*ly), xytext=(px, py),
                arrowprops=dict(arrowstyle="-|>", color="#2EB67D", lw=2.0))
    ax.text(px + 0.85*lx - 0.1, py + 0.85*ly + 0.05, r"$F_{lift}$", color="#2EB67D", fontsize=8)

    # Gravity
    ax.annotate("", xy=(px, py - 0.7), xytext=(px, py),
                arrowprops=dict(arrowstyle="-|>", color="#9B59B6", lw=2.0))
    ax.text(px + 0.05, py - 0.75, "mg", color="#9B59B6", fontsize=8)

    # Flight path angle arc
    angle_arc = np.linspace(0, np.arctan2(dy_v, dx_v), 30)
    ax.plot(px + 0.4*np.cos(angle_arc), py + 0.4*np.sin(angle_arc), color="#888", lw=1.2)
    ax.text(px + 0.5, py - 0.15, r"$\gamma$", fontsize=9, color="#555")

    ax.legend(fontsize=7, loc="lower right")

    # --- Right: front view showing bank angle σ ---
    ax2 = axes[1]
    ax2.set_xlim(-2.5, 2.5)
    ax2.set_ylim(-1.5, 2.5)
    ax2.axis("off")
    ax2.set_facecolor("#F8F9FA")
    ax2.set_title("Front View (σ: bank angle)", fontsize=9, fontweight="bold")

    # Horizon line
    ax2.axhline(0, color="#CCC", lw=1.0, linestyle="--")
    ax2.text(2.1, 0.05, "horiz.", fontsize=7, color="#888")

    # Vehicle dot (coming toward viewer)
    ax2.plot(0, 0.5, "o", ms=12, color="#E05A2B", zorder=5)
    ax2.text(0.15, 0.5, "HGV\n(front)", fontsize=7)

    # Bank angle = 30°
    sigma = np.radians(35)
    # Wing span
    ax2.annotate("", xy=(-1.5*np.cos(sigma), 0.5 - 1.5*np.sin(sigma)),
                xytext=(0, 0.5), arrowprops=dict(arrowstyle="-", color="#555", lw=2.0))
    ax2.annotate("", xy=(1.5*np.cos(sigma), 0.5 + 1.5*np.sin(sigma)),
                xytext=(0, 0.5), arrowprops=dict(arrowstyle="-", color="#555", lw=2.0))

    # Lift direction (perpendicular to wing)
    lx2 = np.sin(sigma); ly2 = np.cos(sigma)
    ax2.annotate("", xy=(0 + 1.0*lx2, 0.5 + 1.0*ly2), xytext=(0, 0.5),
                arrowprops=dict(arrowstyle="-|>", color="#2EB67D", lw=2.0))
    ax2.text(0 + 1.0*lx2 + 0.1, 0.5 + 1.0*ly2, r"$F_{lift}$", color="#2EB67D", fontsize=8)

    # Gravity
    ax2.annotate("", xy=(0, 0.5 - 0.8), xytext=(0, 0.5),
                arrowprops=dict(arrowstyle="-|>", color="#9B59B6", lw=2.0))
    ax2.text(0.05, 0.5 - 0.9, "mg", color="#9B59B6", fontsize=8)

    # Bank angle arc
    bank_arc = np.linspace(np.pi/2, np.pi/2 - sigma, 40)
    r_arc = 0.55
    ax2.plot(0 + r_arc*np.cos(bank_arc), 0.5 + r_arc*np.sin(bank_arc), color="#E05A2B", lw=1.5)
    ax2.text(-0.05, 0.5 + 0.65, r"$\sigma$", fontsize=9, color="#E05A2B", fontweight="bold")

    fig.tight_layout()
    return fig


def _make_hgv_skimming_figure():
    """Illustrate HGV skip-glide oscillation pattern."""
    import numpy as np
    fig, ax = plt.subplots(figsize=(7, 3.5))

    t = np.linspace(0, 1, 500)
    x = t * 1400

    # Boost phase
    t_boost = 0.04
    t_pullup = 0.10
    t_glide_start = 0.15
    t_terminal = 0.88

    # Boost (0 → t_boost): steep ascent
    boost_mask = t < t_boost
    # Pullup: t_boost → t_glide_start
    pullup_mask = (t >= t_boost) & (t < t_glide_start)
    # Glide with skip oscillation
    glide_mask = (t >= t_glide_start) & (t < t_terminal)
    # Terminal
    terminal_mask = t >= t_terminal

    alt = np.zeros_like(t)
    # Boost ascent
    alt[boost_mask] = np.linspace(0, 35, boost_mask.sum())
    # Pullup transition
    alt[pullup_mask] = np.linspace(35, 42, pullup_mask.sum())
    # Skip-glide: slow oscillation around cruise altitude
    t_g = (t[glide_mask] - t_glide_start) / (t_terminal - t_glide_start)
    alt[glide_mask] = 40 + 6 * np.exp(-1.5*t_g) * np.sin(3.5 * np.pi * t_g) - 3 * t_g
    # Terminal dive
    t_term = (t[terminal_mask] - t_terminal) / (1 - t_terminal)
    alt[terminal_mask] = alt[terminal_mask][0] * (1 - t_term**1.5)

    ax.plot(x, alt, color="#2E75B6", lw=2.5)

    # Phase labels
    ax.axvspan(0, t_boost * 1400, alpha=0.15, color="#FF8C00", label="Boost")
    ax.axvspan(t_boost * 1400, t_glide_start * 1400, alpha=0.08, color="#9B59B6",
               label="Pull-up")
    ax.axvspan(t_glide_start * 1400, t_terminal * 1400, alpha=0.08, color="#2E75B6",
               label="Skip-Glide (bank angle control)")
    ax.axvspan(t_terminal * 1400, 1400, alpha=0.15, color="#E05A2B",
               label="Terminal")

    ax.set_xlabel("Horizontal Range (km)", fontsize=10)
    ax.set_ylabel("Altitude (km)", fontsize=10)
    ax.set_title("HGV Skip-Glide Trajectory", fontsize=12, fontweight="bold")
    ax.legend(fontsize=8, loc="upper right")
    ax.set_xlim(0, 1400)
    ax.set_ylim(-3, 58)
    ax.grid(True, alpha=0.3)
    ax.set_facecolor("#F8F9FA")
    fig.tight_layout()
    return fig


def _make_gpu_arch_figure():
    """Illustrate GPU parallelization architecture."""
    import numpy as np
    fig, ax = plt.subplots(figsize=(7.5, 4.0))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6.5)
    ax.axis("off")
    ax.set_facecolor("white")

    from matplotlib.patches import FancyBboxPatch

    def draw_box(cx, cy, w, h, color, text, fontsize=8.5, alpha=0.88):
        rect = FancyBboxPatch((cx - w/2, cy - h/2), w, h,
                              boxstyle="round,pad=0.06",
                              facecolor=color, edgecolor="white",
                              linewidth=1.5, alpha=alpha)
        ax.add_patch(rect)
        ax.text(cx, cy, text, ha="center", va="center",
                fontsize=fontsize, color="white", fontweight="bold",
                linespacing=1.35)

    # CPU section
    draw_box(1.5, 5.5, 2.6, 0.85, "#555555", "CPU (Host)", fontsize=9)

    # GPU outer box
    rect_gpu = FancyBboxPatch((3.2, 0.3), 6.5, 5.8,
                              boxstyle="round,pad=0.1",
                              facecolor="#F0F4FA", edgecolor="#2E75B6",
                              linewidth=2.0, alpha=1.0)
    ax.add_patch(rect_gpu)
    ax.text(6.45, 5.85, "GPU (Device)", ha="center", fontsize=9,
            color="#2E75B6", fontweight="bold")

    # Stream 0: CV UKF
    draw_box(4.4, 4.6, 2.0, 0.8, "#4472C4", "Stream 0\nCV-UKF", fontsize=8)
    # Stream 1: Ballistic UKF
    draw_box(6.45, 4.6, 2.0, 0.8, "#4472C4", "Stream 1\nBallistic-UKF", fontsize=8)
    # Stream 2: CT UKF
    draw_box(8.5, 4.6, 2.0, 0.8, "#4472C4", "Stream 2\nCT-UKF", fontsize=8)

    # Inside each stream: parallel sigma point processing
    for cx_s in [4.4, 6.45, 8.5]:
        draw_box(cx_s - 0.5, 3.55, 0.7, 0.65, "#2E75B6", "19×N\nthreads", fontsize=6.5)
        draw_box(cx_s + 0.5, 3.55, 0.7, 0.65, "#2E75B6", "motion\nmodel", fontsize=6.5)
        ax.annotate("", xy=(cx_s+0.12, 3.55), xytext=(cx_s-0.12, 3.55),
                    arrowprops=dict(arrowstyle="-|>", color="#AAA", lw=1.0))

    ax.text(6.45, 2.85, "Parallel sigma point propagation\n(N_targets x 19 threads per kernel)",
            ha="center", fontsize=7.5, color="#444", style="italic")

    # Sync + IMM combine
    draw_box(6.45, 2.0, 4.5, 0.75, "#2EB67D",
             "Stream sync + IMM weight combination (N_targets threads)", fontsize=7.5)

    # Data association
    draw_box(6.45, 1.05, 4.5, 0.75, "#E05A2B",
             "Cost matrix kernel: N_tracks x N_meas threads", fontsize=7.5)

    # Constant memory
    draw_box(3.75, 2.5, 0.9, 1.2, "#9B59B6", "__const__\nmemory\ng0, Re\nrho0, H", fontsize=6.5)

    # Arrows CPU -> GPU
    ax.annotate("", xy=(3.3, 5.2), xytext=(2.8, 5.2),
                arrowprops=dict(arrowstyle="-|>", color="#555", lw=1.5))
    ax.text(3.05, 5.4, "H2D", fontsize=7.5, color="#555", ha="center")
    ax.annotate("", xy=(2.8, 4.8), xytext=(3.3, 4.8),
                arrowprops=dict(arrowstyle="-|>", color="#555", lw=1.5))
    ax.text(3.05, 4.6, "D2H", fontsize=7.5, color="#555", ha="center")

    fig.tight_layout(rect=[0, 0, 1, 1])
    return fig


def _make_measurement_geometry_figure():
    """Create a 2D diagram showing radar measurement geometry."""
    import numpy as np
    fig, axes = plt.subplots(1, 2, figsize=(7, 3.2))

    # Left: azimuth (top-down view)
    ax = axes[0]
    ax.set_aspect("equal")
    ax.set_xlim(-1.2, 5)
    ax.set_ylim(-1.2, 4.5)
    ax.axis("off")
    ax.set_facecolor("#F8F9FA")
    ax.set_title("Azimuth (Top View)", fontsize=9, fontweight="bold")

    # Radar
    ax.plot(0, 0, "s", ms=9, color="#E05A2B", zorder=5)
    ax.text(0, -0.5, "Radar", ha="center", fontsize=8)

    # Target
    tx, ty = 3.5, 3.0
    ax.plot(tx, ty, "^", ms=9, color="#2E75B6", zorder=5)
    ax.text(tx + 0.15, ty + 0.1, "Target", fontsize=8)

    # Range line
    ax.annotate("", xy=(tx, ty), xytext=(0, 0),
                arrowprops=dict(arrowstyle="-|>", color="#555", lw=1.5))
    ax.text(1.5, 1.9, "r (range)", fontsize=8, color="#555", rotation=40)

    # North reference
    ax.annotate("", xy=(0, 2.5), xytext=(0, 0),
                arrowprops=dict(arrowstyle="-|>", color="#AAA", lw=1.2))
    ax.text(0.1, 2.5, "N", fontsize=8, color="#888")

    # Azimuth arc
    theta_vals = np.linspace(np.pi / 2, np.arctan2(ty, tx), 60)
    ax.plot(1.2 * np.cos(theta_vals), 1.2 * np.sin(theta_vals), color="#2EB67D", lw=1.5)
    ax.text(0.4, 1.4, "θ", fontsize=10, color="#2EB67D", fontweight="bold")

    # Right: elevation (side view)
    ax = axes[1]
    ax.set_aspect("equal")
    ax.set_xlim(-0.5, 5)
    ax.set_ylim(-0.5, 4.0)
    ax.axis("off")
    ax.set_facecolor("#F8F9FA")
    ax.set_title("Elevation (Side View)", fontsize=9, fontweight="bold")

    ax.plot(0, 0, "s", ms=9, color="#E05A2B", zorder=5)
    ax.text(0, -0.35, "Radar", ha="center", fontsize=8)

    tx2, tz2 = 3.5, 2.5
    ax.plot(tx2, tz2, "^", ms=9, color="#2E75B6", zorder=5)
    ax.text(tx2 + 0.15, tz2 + 0.1, "Target", fontsize=8)

    ax.annotate("", xy=(tx2, tz2), xytext=(0, 0),
                arrowprops=dict(arrowstyle="-|>", color="#555", lw=1.5))
    ax.text(1.5, 1.6, "r (range)", fontsize=8, color="#555", rotation=35)

    # Ground reference
    ax.annotate("", xy=(3.5, 0), xytext=(0, 0),
                arrowprops=dict(arrowstyle="-|>", color="#AAA", lw=1.2))
    ax.text(3.5, -0.3, "Horiz.", fontsize=8, color="#888")

    # Elevation arc
    phi_vals = np.linspace(0, np.arctan2(tz2, tx2), 60)
    ax.plot(1.2 * np.cos(phi_vals), 1.2 * np.sin(phi_vals), color="#9B59B6", lw=1.5)
    ax.text(1.3, 0.5, "φ", fontsize=10, color="#9B59B6", fontweight="bold")

    # Altitude annotation
    ax.plot([tx2, tx2], [0, tz2], color="#CCC", lw=1.0, linestyle="--")
    ax.text(tx2 + 0.1, tz2 / 2, "z", fontsize=9, color="#888")

    fig.tight_layout()
    return fig


def add_section_target_generator(doc):
    """Section 2: Target Generator."""
    doc.add_heading("2. ターゲットジェネレータ", level=1)

    add_paragraph_with_font(
        doc,
        "ターゲットジェネレータは、弾道ミサイル・HGV・巡航ミサイル等の目標軌道を"
        "物理モデルに基づき生成するコンポーネントである。"
        "生成された軌道はレーダシミュレータへの入力として使用される。",
        size=10,
    )

    # --- 2.1 Overview ---
    doc.add_heading("2.1 軌道タイプ", level=2)
    add_paragraph_with_font(
        doc,
        "サポートする5種類の運動モデルを表2-1に示す。"
        "各軌道タイプの概念図を図2-1に示す。",
        size=10,
    )

    add_table_caption(doc, "表2-1  軌道タイプ一覧")
    add_table_with_header(
        doc,
        ["モデル名", "略称", "特徴", "代表的用途"],
        [
            ("CONSTANT_VELOCITY", "CV", "等速直線飛翔", "巡航ミサイル・慣性飛翔"),
            ("CONSTANT_ACCELERATION", "CA", "等加速度直線飛翔", "加速段・減速段の近似"),
            ("MANEUVERING", "MN", "ランダム機動（連続旋回）", "機動目標シミュレーション"),
            ("BALLISTIC_MISSILE", "BM", "物理ベース弾道軌道（RK4）", "弾道ミサイル（ICBM等）"),
            ("HYPERSONIC_GLIDE", "HGV", "滑空飛翔（揚力・抗力）", "極超音速滑空体"),
        ],
        col_widths=[4.5, 2.0, 4.5, 5.0],
    )

    # Figure: trajectory types
    doc.add_paragraph()
    fig = _make_trajectory_figure()
    fig_path = save_figure(fig, "traj_types.png")
    add_figure_to_doc(doc, fig_path,
                      caption="図2-1  各軌道タイプの概念図（水平距離 vs 高度）",
                      width_inches=5.5)

    # --- 2.2 State vector ---
    doc.add_heading("2.2 目標状態ベクトル", level=2)
    add_paragraph_with_font(
        doc,
        "ターゲットジェネレータが保持する状態ベクトルは9次元であり、"
        "位置・速度・加速度を含む。座標系はセンサ中心の直交座標系（x=東, y=北, z=高度）を使用する。",
        size=10,
    )
    add_math_image(
        doc,
        r"\mathbf{x}_{gt} = [x,\ y,\ z,\ v_x,\ v_y,\ v_z,\ a_x,\ a_y,\ a_z]^T",
        fontsize=16,
    )
    add_symbol_legend(doc, [
        ("x_{gt}", "目標の状態ベクトル（9次元）"),
        ("x, y, z", "位置（東・北・高度） [m]"),
        ("v_x, v_y, v_z", "速度成分 [m/s]"),
        ("a_x, a_y, a_z", "加速度成分 [m/s²]"),
        ("T (上付き)", "転置（列ベクトルを表す）"),
    ])

    # --- 2.3 Ballistic missile model ---
    doc.add_heading("2.3 弾道ミサイルモデル（RK4）", level=2)
    add_paragraph_with_font(
        doc,
        "弾道ミサイルはブースト・中間飛翔・終端の3フェーズで構成される（図2-2）。"
        "中間飛翔フェーズでは重力と大気抵抗を考慮した物理方程式を"
        "4次ルンゲ・クッタ法（RK4）で数値積分する。",
        size=10,
    )

    # Figure: ballistic phases
    fig2 = _make_ballistic_phases_figure()
    fig2_path = save_figure(fig2, "ballistic_phases.png")
    add_figure_to_doc(doc, fig2_path,
                      caption="図2-2  弾道ミサイル飛翔フェーズ",
                      width_inches=5.5)

    doc.add_paragraph()
    add_paragraph_with_font(doc, "RK4数値積分（各フレーム）:", size=10, bold=True)
    add_math_image(
        doc,
        r"\mathbf{k}_1 = f(\mathbf{x}_n)",
        fontsize=14,
    )
    add_math_image(
        doc,
        r"\mathbf{k}_2 = f\left(\mathbf{x}_n + \frac{\Delta t}{2}\mathbf{k}_1\right)",
        fontsize=14,
    )
    add_math_image(
        doc,
        r"\mathbf{k}_3 = f\left(\mathbf{x}_n + \frac{\Delta t}{2}\mathbf{k}_2\right)",
        fontsize=14,
    )
    add_math_image(
        doc,
        r"\mathbf{k}_4 = f(\mathbf{x}_n + \Delta t \, \mathbf{k}_3)",
        fontsize=14,
    )
    add_math_image(
        doc,
        r"\mathbf{x}_{n+1} = \mathbf{x}_n + \frac{\Delta t}{6}"
        r"(\mathbf{k}_1 + 2\mathbf{k}_2 + 2\mathbf{k}_3 + \mathbf{k}_4)",
        fontsize=15,
    )
    add_paragraph_with_font(
        doc,
        "ここで f は重力・大気抵抗を含む運動方程式（セクション7.2参照）である。"
        "軌道はシミュレーション開始時に50msキャッシュで事前計算され、"
        "getPosition(t)呼び出しは補間で高速に返す。",
        size=10,
    )
    add_symbol_legend(doc, [
        ("k_1, k_2, k_3, k_4", "RK4の各段の勾配ベクトル"),
        ("f(x)", "運動方程式（重力＋大気抵抗）"),
        ("x_n", "タイムステップ n における状態ベクトル"),
        ("x_{n+1}", "次ステップの状態ベクトル"),
        ("Delta t", "積分タイムステップ [s]"),
    ])

    # --- 2.4 HGV ---
    doc.add_heading("2.4 HGV（極超音速滑空体）モデル", level=2)
    add_paragraph_with_font(
        doc,
        "HGV（Hypersonic Glide Vehicle）は、ロケットブーストで大気圏外へ射出された後、"
        "大気圏上層で揚力を利用してスキップ・グライドしながら目標へ向かう。"
        "本システムのHGVモデルは、揚力・抗力・重力を考慮した完全3次元運動方程式を"
        "4次ルンゲ・クッタ法で数値積分し、フェーズごとにバンク角を制御する。",
        size=10,
    )

    # 2.4.1 飛翔フェーズ
    doc.add_heading("2.4.1 飛翔フェーズ", level=3)
    add_paragraph_with_font(
        doc,
        "HGVの飛翔は5つのフェーズに分けられる。各フェーズの特性と遷移条件を表2-3に示す。",
        size=10,
    )

    add_table_caption(doc, "表2-3  HGV飛翔フェーズ一覧")
    add_table_with_header(
        doc,
        ["フェーズ", "状態", "特性", "遷移条件"],
        [
            ("BOOST", "ブースト上昇",
             "推力 F_thrust を鉛直・水平に配分。\n重力ターン制御。",
             "燃焼時間 t_boost が終了"),
            ("PULLUP", "プルアップ",
             "ブースト終了後の姿勢回復。\nバンク角で揚力を利用し仰角を 0 へ収束。",
             "|γ| < 0.08 rad かつ L/W > 0.5\nかつ経過時間 > 10 s"),
            ("GLIDE", "スキップ滑空",
             "揚力・抗力・重力の3力で軌道制御。\nバンク角で横方向誘導。\n大気密度変動によりスキップ振動が生じる。",
             "目標までの距離 < 15～30 km"),
            ("TERMINAL", "終末突入",
             "急降下で目標に向かう。\n0.3 rad のサイン波型機動を重畳。",
             "目標到達または高度 < 1 km"),
            ("MIDCOURSE", "弾道中間",
             "推力・揚力なし。重力・抗力のみ（弾道飛翔）。",
             "—（非 HGV 用）"),
        ],
        col_widths=[2.2, 2.5, 5.8, 4.5],
    )

    # 2.4.2 運動方程式
    doc.add_heading("2.4.2 運動方程式（3次元）", level=3)
    add_paragraph_with_font(
        doc,
        "状態ベクトルは位置と速度の6成分 [x, y, z, vx, vy, vz] である。"
        "重力加速度 g(z)、大気密度 ρ(z) はいずれも高度に依存する。"
        "HGVに作用する力と幾何学的関係を図2-3に示す。",
        size=10,
    )

    # Figure: HGV forces
    fig_forces = _make_hgv_forces_figure()
    fig_forces_path = save_figure(fig_forces, "hgv_forces.png")
    add_figure_to_doc(
        doc, fig_forces_path,
        caption="図2-3  HGVに作用する力（左: 飛翔経路面, 右: バンク角）",
        width_inches=6.5,
    )

    add_paragraph_with_font(doc, "位置微分:", size=10, bold=True)
    add_math_image(doc,
        r"\dot{x} = v_x, \quad \dot{y} = v_y, \quad \dot{z} = v_z",
        fontsize=14)

    add_paragraph_with_font(doc, "速度微分:", size=10, bold=True)
    add_math_image(doc,
        r"\dot{v}_x = \frac{F_{D,x} + F_{L,x}}{m}, \quad"
        r"\dot{v}_y = \frac{F_{D,y} + F_{L,y}}{m}, \quad"
        r"\dot{v}_z = \frac{F_{D,z} + F_{L,z}}{m} - g(z)",
        fontsize=13)

    add_paragraph_with_font(doc, "高度依存の重力と大気密度:", size=10, bold=True)
    add_math_image(doc,
        r"g(z) = g_0 \left(\frac{R_E}{R_E + z}\right)^2, \quad"
        r"\rho(z) = \rho_0 \exp\left(-\frac{z}{H}\right)",
        fontsize=14)
    add_paragraph_with_font(
        doc,
        "ここで g₀ = 9.80665 m/s²、R_E = 6,371,000 m（地球半径）、"
        "ρ₀ = 1.225 kg/m³（海面密度）、H = 8,500 m（スケール高度）。",
        size=10,
    )
    add_symbol_legend(doc, [
        ("dot{x}, dot{y}, dot{z}", "位置の時間微分（速度成分） [m/s]"),
        ("v_x, v_y, v_z", "速度成分 [m/s]"),
        ("dot{v}_x, dot{v}_y, dot{v}_z", "速度の時間微分（加速度成分） [m/s²]"),
        ("F_{D,x/y/z}", "抗力ベクトルの各成分 [N]"),
        ("F_{L,x/y/z}", "揚力ベクトルの各成分 [N]"),
        ("m", "機体質量 [kg]"),
        ("g(z)", "高度 z における重力加速度 [m/s²]"),
        ("g_0", "標準重力加速度 = 9.80665 m/s²"),
        ("R_E", "地球半径 = 6,371,000 m"),
        ("rho(z)", "高度 z における大気密度 [kg/m³]"),
        ("rho_0", "海面大気密度 = 1.225 kg/m³"),
        ("H", "大気スケール高度 = 8,500 m"),
    ])

    # 2.4.3 揚力・抗力モデル
    doc.add_heading("2.4.3 抗力と揚力", level=3)
    add_paragraph_with_font(doc, "抗力（速度方向逆向き）:", size=10, bold=True)
    add_math_image(doc,
        r"\mathbf{F}_{drag} = -\frac{1}{2}\rho(z)\,v^2\,C_D\,A\;\hat{v}",
        fontsize=14)
    add_paragraph_with_font(doc, "揚力の大きさ:", size=10, bold=True)
    add_math_image(doc,
        r"L = \frac{1}{2}\rho(z)\,v^2\,C_L\,A, \quad C_L = C_D \cdot (L/D)",
        fontsize=14)
    add_paragraph_with_font(
        doc,
        "揚力方向は速度ベクトル v̂ に垂直な2方向から合成される。"
        "まず鉛直面内の上向き単位ベクトル n̂_up を次式で定義する。",
        size=10,
    )
    add_math_image(doc,
        r"\hat{n}_{up} = \frac{\hat{z} - (\hat{z}\cdot\hat{v})\hat{v}}{|\hat{z} - (\hat{z}\cdot\hat{v})\hat{v}|}",
        fontsize=14)
    add_paragraph_with_font(
        doc,
        "横方向単位ベクトル n̂_lat は次式で得られる。",
        size=10,
    )
    add_math_image(doc,
        r"\hat{n}_{lat} = \hat{v} \times \hat{n}_{up}",
        fontsize=14)
    add_paragraph_with_font(doc, "揚力ベクトル（バンク角 σ を反映）:", size=10, bold=True)
    add_math_image(doc,
        r"\mathbf{F}_{lift} = L\,(\cos\sigma\;\hat{n}_{up} + \sin\sigma\;\hat{n}_{lat})",
        fontsize=14)
    add_symbol_legend(doc, [
        ("F_{drag}", "抗力ベクトル（速度逆方向） [N]"),
        ("rho(z)", "高度依存大気密度 [kg/m³]"),
        ("v", "速度の大きさ [m/s]"),
        ("C_D", "抗力係数 [-]"),
        ("A", "代表断面積 [m²]"),
        ("hat{v}", "速度方向の単位ベクトル"),
        ("L", "揚力の大きさ [N]"),
        ("C_L", "揚力係数（= C_D × L/D）[-]"),
        ("L/D", "揚抗比（= 2.0）[-]"),
        ("hat{n}_{up}", "鉛直面内の上向き単位ベクトル"),
        ("hat{z}", "鉛直（高度）方向の単位ベクトル"),
        ("hat{n}_{lat}", "横方向単位ベクトル（= v × n_up）"),
        ("F_{lift}", "揚力ベクトル [N]"),
        ("sigma", "バンク角（機体ロール角） [rad]"),
    ])

    # 2.4.4 バンク角制御
    doc.add_heading("2.4.4 バンク角制御則", level=3)
    add_paragraph_with_font(
        doc,
        "バンク角 σ はフェーズごとに異なる制御則で決定される。"
        "γ は飛翔経路角（仰角）、Δψ は目標方位との偏差、"
        "K_hdg = 2.0 rad⁻¹ は方位誘導ゲイン、L/W は揚力／重力比、"
        "f は機動周波数、amp は機動振幅（0.3 rad）。",
        size=10,
    )
    add_table_caption(doc, "表2-4  フェーズ別バンク角制御則")
    add_table_with_header(
        doc,
        ["フェーズ", "バンク角 σ の計算式", "説明"],
        [
            ("PULLUP",
             "cos(σ) = mg·cos(γ) / L",
             "揚力水平分力で仰角をゼロへ回復"),
            ("GLIDE",
             "σ = −K_hdg · Δψ",
             "横方向バンクで方位偏差を修正（K_hdg = 2.0 rad⁻¹）"),
            ("TERMINAL",
             "σ = π/2 + σ_hdg + 0.3·amp·sin(2π·f·t)",
             "方位誘導＋サイン波機動で被撃墜率を低減"),
        ],
        col_widths=[2.5, 6.0, 7.5],
    )

    # 2.4.5 主要パラメータ
    doc.add_heading("2.4.5 主要パラメータ", level=3)
    add_paragraph_with_font(
        doc,
        "HGVモデルの主要パラメータを表2-5に示す。"
        "スキップ滑空の典型的な高度プロファイルを図2-4に示す。",
        size=10,
    )
    add_table_caption(doc, "表2-5  HGVモデルパラメータ")
    add_table_with_header(
        doc,
        ["パラメータ", "記号", "代表値", "説明"],
        [
            ("抗力係数×面積", "C_D·A", "0.5 m²", "抗力計算に使用する有効面積"),
            ("揚抗比", "L/D", "2.0", "揚力係数 / 抗力係数"),
            ("質量", "m", "1,000 kg", "機体質量"),
            ("ブースト推力", "F_thrust", "200,000 N", "ブーストフェーズの推力"),
            ("燃焼時間", "t_boost", "60 s", "ブーストフェーズ継続時間"),
            ("方位誘導ゲイン", "K_hdg", "2.0 rad⁻¹", "GLIDEフェーズの横方向制御ゲイン"),
            ("終末機動振幅", "amp", "0.3 rad", "TERMINALフェーズのバンク角振幅"),
            ("終末機動周波数", "f", "0.1 Hz", "TERMINALフェーズのサイン波周波数"),
        ],
        col_widths=[4.5, 2.5, 3.0, 6.0],
    )

    # Figure: skip-glide
    doc.add_paragraph()
    fig_skim = _make_hgv_skimming_figure()
    fig_skim_path = save_figure(fig_skim, "hgv_skimming.png")
    add_figure_to_doc(
        doc, fig_skim_path,
        caption="図2-4  HGVスキップ滑空の典型的な高度プロファイル",
        width_inches=6.0,
    )

    doc.add_page_break()


def add_section_radar_simulator(doc):
    """Section 3: Radar Simulator."""
    doc.add_heading("3. レーダシミュレータ", level=1)

    add_paragraph_with_font(
        doc,
        "レーダシミュレータは、実際のレーダの観測特性を模擬するコンポーネントである。"
        "Swerling IIモデルによる検出確率変動、測定ノイズ、クラッタを生成し、"
        "追尾アルゴリズムへのリアルな入力を提供する。",
        size=10,
    )

    # --- 3.1 Measurement model ---
    doc.add_heading("3.1 観測モデル", level=2)
    add_paragraph_with_font(
        doc,
        "レーダはセンサ位置 (x_s, y_s, z_s) から目標位置 (x, y, z) に対して"
        "4次元観測ベクトルを生成する。各観測量の定義を表3-1に示す。"
        "観測ジオメトリを図3-1に示す。",
        size=10,
    )

    add_table_caption(doc, "表3-1  観測量の定義")
    add_table_with_header(
        doc,
        ["観測量", "記号", "定義式", "単位"],
        [
            ("距離（スラントレンジ）", "r",
             "√(Δx²+Δy²+Δz²)", "m"),
            ("方位角", "θ",
             "atan2(Δy, Δx)", "rad"),
            ("仰角", "φ",
             "atan2(Δz, r_horiz)", "rad"),
            ("ドップラー速度", "ḋ",
             "(Δx·vx+Δy·vy+Δz·vz)/r", "m/s"),
        ],
        col_widths=[4.5, 2.0, 5.5, 2.0],
    )

    # Figure: measurement geometry
    doc.add_paragraph()
    fig_geom = _make_measurement_geometry_figure()
    fig_geom_path = save_figure(fig_geom, "meas_geometry.png")
    add_figure_to_doc(doc, fig_geom_path,
                      caption="図3-1  レーダ観測ジオメトリ（左: 水平面, 右: 垂直面）",
                      width_inches=5.5)

    doc.add_paragraph()
    add_paragraph_with_font(
        doc,
        "各観測量には独立なガウスノイズが付加される（観測ノイズ共分散行列 R の詳細はセクション4.3参照）。",
        size=10,
    )

    # --- 3.2 Swerling II ---
    doc.add_heading("3.2 Swerling II 検出モデル", level=2)
    add_paragraph_with_font(
        doc,
        "本システムはSwerling IIモデルを採用する。"
        "Swerling IIモデルでは目標のレーダ断面積（RCS）がパルス間でランダムに変動し、"
        "指数分布に従う。",
        size=10,
    )

    add_paragraph_with_font(doc, "SNRの距離依存性（レーダ方程式）:", size=10, bold=True)
    add_math_image(
        doc,
        r"\mathrm{SNR}_{avg}(r) = \mathrm{SNR}_{ref} - 40 \log_{10}\!\left(\frac{r}{1000}\right)"
        r"\quad [\mathrm{dB}]",
        fontsize=14,
    )
    add_paragraph_with_font(
        doc,
        "ここで SNR_ref は基準距離 1 km における平均 SNR (dB) であり、"
        "ユーザが指定した基準距離 R_ref における検出確率 P_D から自動導出される。",
        size=10,
    )

    add_paragraph_with_font(doc, "SNR_ref の自動導出:", size=10, bold=True)
    add_math_image(
        doc,
        r"\gamma_T = -\ln(P_{FA})",
        fontsize=14,
    )
    add_math_image(
        doc,
        r"\mathrm{SNR}_{avg}^{lin}(R_{ref}) = \frac{\gamma_T}{-\ln(P_D^{ref})}",
        fontsize=14,
    )
    add_math_image(
        doc,
        r"\mathrm{SNR}_{ref} = 10\log_{10}\!\left(\mathrm{SNR}_{avg}^{lin}\right)"
        r"+ 40\log_{10}\!\left(\frac{R_{ref}}{1000}\right)",
        fontsize=14,
    )

    add_paragraph_with_font(doc, "パルスごとの瞬時SNR:", size=10, bold=True)
    add_math_image(
        doc,
        r"\mathrm{SNR}_{inst} = \mathrm{SNR}_{avg} \cdot \xi, \quad \xi \sim \mathrm{Exp}(1)",
        fontsize=14,
    )

    add_paragraph_with_font(doc, "CFAR（一定誤警報率）検出判定（SNR_inst ≥ γ_T のとき検出）:", size=10, bold=True)
    add_math_image(
        doc,
        r"\mathrm{SNR}_{inst} \geq \gamma_T = -\ln(P_{FA})",
        fontsize=14,
    )

    add_paragraph_with_font(doc, "パルス当たり検出確率（理論値）:", size=10, bold=True)
    add_math_image(
        doc,
        r"P_D = \exp\!\left(\frac{-\gamma_T}{\mathrm{SNR}_{avg}^{lin}}\right)",
        fontsize=14,
    )
    add_paragraph_with_font(
        doc,
        "各P_FA値に対するP_D vs SNRの関係を図3-2に示す。",
        size=10,
    )
    add_symbol_legend(doc, [
        ("SNR_{avg}(r)", "距離 r における平均SNR [dB]"),
        ("SNR_{ref}", "基準距離1 km での平均SNR [dB]"),
        ("r", "目標までのスラントレンジ [m]"),
        ("log_{10}", "常用対数"),
        ("gamma_T", "CFAR検出閾値（= -ln P_FA）"),
        ("P_{FA}", "誤警報確率 [-]"),
        ("ln", "自然対数"),
        ("SNR_{avg}^{lin}", "線形スケールの平均SNR"),
        ("R_{ref}", "ユーザ設定の基準距離 [m]"),
        ("P_D^{ref}", "基準距離における所望検出確率 [-]"),
        ("SNR_{inst}", "パルスごとの瞬時SNR（ランダム変動）"),
        ("xi", "指数分布乱数（RCS変動を模擬）"),
        ("Exp(1)", "平均1の指数分布"),
        ("P_D", "パルス当たりの検出確率 [-]"),
    ])

    # Figure: P(D) vs SNR
    fig_pd = _make_swerling_pd_figure()
    fig_pd_path = save_figure(fig_pd, "swerling_pd.png")
    add_figure_to_doc(doc, fig_pd_path,
                      caption="図3-2  Swerling II モデル: 検出確率 P(D) vs 平均SNR",
                      width_inches=5.0)

    # --- 3.3 Detection chain ---
    doc.add_heading("3.3 検出処理チェーン", level=2)
    add_paragraph_with_font(
        doc,
        "各フレームにおける検出処理は図3-3に示すチェーンで実行される。"
        "各ステップの詳細を表3-2に示す。",
        size=10,
    )

    # Figure: detection chain
    fig_chain = _make_detection_chain_figure()
    fig_chain_path = save_figure(fig_chain, "detection_chain.png")
    add_figure_to_doc(doc, fig_chain_path,
                      caption="図3-3  Swerling II 検出処理チェーン",
                      width_inches=6.0)

    doc.add_paragraph()
    add_table_caption(doc, "表3-2  検出処理チェーン各ステップの詳細")
    add_table_with_header(
        doc,
        ["ステップ", "処理内容", "判定条件"],
        [
            ("①  FOV判定", "センサの視野角（最大仰角）確認",
             "φ ≤ max_elevation_angle"),
            ("②  ビーム方向判定", "ビームステアリング方向との角度差確認",
             "Δangle ≤ beam_width / 2"),
            ("③  SNR計算", "距離・SNR_ref からSNR_avg(r) 算出",
             "SNR_avg = SNR_ref − 40·log10(r/1km) [dB]"),
            ("④  指数乱数生成", "パルスごとのRCS変動を模擬",
             "ξ ~ Exp(1)"),
            ("⑤  CFAR閾値比較", "SNR_inst と CFAR 閾値を比較",
             "SNR_inst ≥ γ_T = −ln(P_FA)"),
        ],
        col_widths=[2.5, 5.5, 8.0],
    )

    # --- 3.4 Clutter ---
    doc.add_heading("3.4 クラッタモデル", level=2)
    add_paragraph_with_font(
        doc,
        "クラッタ（偽観測）はポアソン過程でモデル化される。"
        "各フレームに生成されるクラッタ点数は期待値 λ_c のポアソン分布に従い、"
        "各クラッタはレーダの最大レンジ・FOV内に一様分布する。"
        "クラッタモデルのパラメータを表3-3に示す。",
        size=10,
    )

    add_math_image(
        doc,
        r"N_{clutter} \sim \mathrm{Poisson}(\lambda_c)",
        fontsize=14,
    )
    add_symbol_legend(doc, [
        ("N_{clutter}", "1フレームに生成されるクラッタ点数"),
        ("lambda_c", "クラッタ発生期待点数（ポアソン分布の平均） [点/フレーム]"),
        ("Poisson(lambda_c)", "期待値 λ_c のポアソン分布"),
    ])

    add_table_caption(doc, "表3-3  クラッタモデルパラメータ")
    add_table_with_header(
        doc,
        ["パラメータ", "記号", "デフォルト値", "説明"],
        [
            ("クラッタ期待点数", "λ_c", "5 点/フレーム", "ポアソン分布の期待値"),
            ("クラッタ距離範囲", "r_max", "max_range_m", "最大観測レンジ"),
            ("クラッタ仰角範囲", "φ_max", "max_elevation_angle", "センサFOV最大仰角"),
        ],
        col_widths=[4.0, 2.5, 3.5, 6.0],
    )

    doc.add_page_break()


def add_section_2(doc):
    """Section 4: State-Space Model (renumbered)."""
    doc.add_heading("4. 状態空間モデル", level=1)

    # --- State vector ---
    doc.add_heading("4.1 状態ベクトル", level=2)
    add_paragraph_with_font(
        doc,
        "状態ベクトルは9次元（STATE_DIM=9）で、位置・速度・加速度を含む（表4-1）。",
        size=10,
    )
    add_math_image(doc, r"\mathbf{x} = [x, y, z, v_x, v_y, v_z, a_x, a_y, a_z]^T", fontsize=16)

    add_table_caption(doc, "表4-1  状態変数の定義")
    add_table_with_header(
        doc,
        ["状態変数", "物理量", "単位", "説明"],
        [
            ("x, y, z", "位置", "m", "目標原点座標系（x=東, y=北, z=高度）"),
            ("vx, vy, vz", "速度", "m/s", "各軸方向の速度成分"),
            ("ax, ay, az", "加速度", "m/s²", "各軸方向の加速度成分"),
        ],
        col_widths=[3.0, 2.5, 2.0, 8.5],
    )

    doc.add_paragraph()

    # --- Measurement vector ---
    doc.add_heading("4.2 観測ベクトル", level=2)
    add_paragraph_with_font(
        doc,
        "観測ベクトルは4次元（MEAS_DIM=4）で、レーダ観測量を表す（表4-2）。",
        size=10,
    )
    add_math_image(doc, r"\mathbf{z} = [r, \theta, \phi, \dot{d}]^T", fontsize=16)

    add_table_caption(doc, "表4-2  観測変数の定義")
    add_table_with_header(
        doc,
        ["観測変数", "物理量", "単位", "定義"],
        [
            ("r", "距離", "m", "レーダから目標までのスラントレンジ"),
            ("θ", "方位角", "rad", "atan2(dy, dx)"),
            ("φ", "仰角", "rad", "atan2(dz, r_horiz)"),
            ("ḋ", "ドップラー速度", "m/s", "(dx·vx + dy·vy + dz·vz) / r"),
        ],
        col_widths=[2.5, 3.0, 2.0, 8.5],
    )

    doc.add_paragraph()
    add_paragraph_with_font(
        doc,
        "ここで、dx = x - sensor_x, dy = y - sensor_y, dz = z - sensor_z "
        "はセンサ位置からの相対座標である。",
        size=10,
    )

    # --- Measurement noise ---
    doc.add_heading("4.3 観測ノイズ", level=2)
    add_paragraph_with_font(
        doc,
        "観測ノイズ共分散行列 R はデフォルトで以下の対角行列で定義される（表4-3）。",
        size=10,
    )

    add_table_caption(doc, "表4-3  観測ノイズパラメータ")
    add_table_with_header(
        doc,
        ["観測量", "パラメータ", "デフォルト値", "備考"],
        [
            ("距離", "σ_r", "10.0 m", ""),
            ("方位角", "σ_θ", "0.01 rad", "≈ 0.57°"),
            ("仰角", "σ_φ", "0.01 rad", "≈ 0.57°"),
            ("ドップラー", "σ_ḋ", "2.0 m/s", ""),
        ],
        col_widths=[3.0, 3.0, 3.5, 3.5],
    )

    doc.add_paragraph()
    add_math_image(
        doc,
        r"R = \mathrm{diag}(\sigma_r^2,\ \sigma_\theta^2,\ \sigma_\phi^2,\ \sigma_{\dot{d}}^2)",
        fontsize=15,
    )
    add_symbol_legend(doc, [
        ("R", "観測ノイズ共分散行列（4×4対角行列）"),
        ("diag(...)", "対角行列コンストラクタ"),
        ("sigma_r", "距離観測の標準偏差（= 10 m）"),
        ("sigma_theta", "方位角観測の標準偏差（= 0.01 rad）"),
        ("sigma_phi", "仰角観測の標準偏差（= 0.01 rad）"),
        ("sigma_{dot{d}}", "ドップラー観測の標準偏差（= 2 m/s）"),
    ])

    doc.add_page_break()


def add_section_3(doc):
    """Section 3: UKF."""
    doc.add_heading("5. Unscented Kalman Filter (UKF)", level=1)

    add_paragraph_with_font(
        doc,
        "Unscented Kalman Filter（UKF）は、非線形系に対するカルマンフィルタの拡張であり、"
        "シグマポイントによる統計的線形化を用いる。"
        "本システムではCUDAによりシグマポイント演算をGPU上で並列処理する。",
        size=10,
    )

    # --- Parameters ---
    doc.add_heading("5.1 UKFパラメータ", level=2)
    add_paragraph_with_font(doc, "UKFの主要パラメータを表5-1に示す。", size=10)
    add_table_caption(doc, "表5-1  UKFパラメータ")
    add_table_with_header(
        doc,
        ["パラメータ", "記号", "値", "説明"],
        [
            ("拡散パラメータ", "α", "0.5", "シグマポイントの拡散度を制御"),
            ("分布パラメータ", "β", "2.0", "ガウス分布では β=2 が最適"),
            ("二次スケーリング", "κ", "0.0", "通常 0 に設定"),
            ("状態次元", "n", "9", "STATE_DIM"),
            ("スケーリング", "λ", "α²(n+κ)−n", "= 0.5²×9 − 9 = −6.75"),
        ],
        col_widths=[3.5, 2.0, 3.5, 7.0],
    )

    doc.add_paragraph()
    add_paragraph_with_font(doc, "λの計算式:", size=10, bold=True)
    add_math_image(doc, r"\lambda = \alpha^2 (n + \kappa) - n", fontsize=16)
    add_symbol_legend(doc, [
        ("lambda", "UKFスケーリングパラメータ（= α²(n+κ)−n）"),
        ("alpha", "シグマポイント拡散パラメータ（= 0.5）"),
        ("n", "状態次元数（= 9）"),
        ("kappa", "二次スケーリングパラメータ（= 0.0）"),
        ("beta", "分布パラメータ（ガウス分布では 2.0 が最適）"),
    ])

    # --- Sigma points ---
    doc.add_heading("5.2 シグマポイント", level=2)
    add_paragraph_with_font(
        doc,
        "2n+1 = 19個のシグマポイントを以下のように生成する（√(·) はコレスキー分解を表す）。",
        size=10,
    )

    add_math_image(doc, r"\chi_0 = \bar{\mathbf{x}}", fontsize=15)
    add_math_image(
        doc,
        r"\chi_i = \bar{\mathbf{x}} + \left(\sqrt{(n+\lambda)\mathbf{P}}\right)_i "
        r"\quad (i=1,\ldots,n)",
        fontsize=15,
    )
    add_math_image(
        doc,
        r"\chi_{i+n} = \bar{\mathbf{x}} - \left(\sqrt{(n+\lambda)\mathbf{P}}\right)_i "
        r"\quad (i=1,\ldots,n)",
        fontsize=15,
    )
    add_symbol_legend(doc, [
        ("chi_0", "中心シグマポイント（= 状態推定値）"),
        ("chi_i  (i=1..n)", "正側シグマポイント"),
        ("chi_{i+n} (i=1..n)", "負側シグマポイント"),
        ("bar{x}", "現在の状態推定値（平均）"),
        ("sqrt{(n+lambda)P}_i", "行列の平方根（コレスキー分解）のi列目"),
        ("P", "現在の推定誤差共分散行列"),
        ("n", "状態次元数（= 9）、シグマポイント総数 = 2n+1 = 19"),
        ("lambda", "スケーリングパラメータ（= −6.75）"),
    ])

    # --- Weights ---
    doc.add_heading("5.3 重み係数", level=2)
    add_paragraph_with_font(doc, "平均用重み:", size=10, bold=True)
    add_math_image(doc, r"W_0^m = \frac{\lambda}{n + \lambda}", fontsize=15)
    add_math_image(
        doc,
        r"W_i^m = \frac{1}{2(n + \lambda)} \quad (i=1,\ldots,2n)",
        fontsize=15,
    )

    add_paragraph_with_font(doc, "共分散用重み:", size=10, bold=True)
    add_math_image(
        doc,
        r"W_0^c = \frac{\lambda}{n + \lambda} + (1 - \alpha^2 + \beta)",
        fontsize=15,
    )
    add_math_image(
        doc,
        r"W_i^c = \frac{1}{2(n + \lambda)} \quad (i=1,\ldots,2n)",
        fontsize=15,
    )
    add_symbol_legend(doc, [
        ("W_0^m", "中心シグマポイントの平均用重み"),
        ("W_i^m  (i=1..2n)", "正負シグマポイントの平均用重み（各 1/(2(n+λ))）"),
        ("W_0^c", "中心シグマポイントの共分散用重み（ベータ補正を含む）"),
        ("W_i^c  (i=1..2n)", "正負シグマポイントの共分散用重み（= W_i^m と同値）"),
        ("beta", "分布パラメータ（= 2.0）。共分散重みの尖度補正に使用"),
        ("lambda", "スケーリングパラメータ"),
        ("n", "状態次元数（= 9）"),
    ])

    # --- Prediction ---
    doc.add_heading("5.4 予測ステップ", level=2)
    add_paragraph_with_font(doc, "予測ステップは以下の手順で実行する。", size=10)

    add_paragraph_with_font(doc, "手順1: シグマポイント生成", size=10, bold=True)
    add_paragraph_with_font(
        doc,
        "現在の状態推定値 x̂ と共分散行列 P から、上記のシグマポイント χᵢ を生成する。",
        size=10,
    )

    add_paragraph_with_font(doc, "手順2: 運動モデルによる伝播", size=10, bold=True)
    add_math_image(doc, r"\chi_i^* = f(\chi_i, \Delta t)", fontsize=15)

    add_paragraph_with_font(doc, "手順3: 予測平均", size=10, bold=True)
    add_math_image(
        doc,
        r"\hat{\mathbf{x}}^- = \sum_{i=0}^{2n} W_i^m \chi_i^*",
        fontsize=15,
    )

    add_paragraph_with_font(doc, "手順4: 予測共分散", size=10, bold=True)
    add_math_image(
        doc,
        r"\mathbf{P}^- = \sum_{i=0}^{2n} W_i^c "
        r"(\chi_i^* - \hat{\mathbf{x}}^-)(\chi_i^* - \hat{\mathbf{x}}^-)^T + \mathbf{Q} \cdot \Delta t",
        fontsize=15,
    )
    add_symbol_legend(doc, [
        ("chi_i^*", "運動モデル適用後の伝播済みシグマポイント"),
        ("f(chi_i, Delta t)", "運動モデル関数（CV / 弾道 / CT）"),
        ("Delta t", "タイムステップ [s]"),
        ("hat{x}^-", "予測状態推定値（更新前）"),
        ("W_i^m", "平均用重み係数"),
        ("P^-", "予測誤差共分散行列"),
        ("W_i^c", "共分散用重み係数"),
        ("Q", "プロセスノイズ共分散行列"),
        ("T (上付き)", "行列転置"),
    ])

    # --- Update ---
    doc.add_heading("5.5 更新ステップ", level=2)

    add_paragraph_with_font(doc, "手順1: 観測空間への変換", size=10, bold=True)
    add_math_image(doc, r"\mathbf{Z}_i = h(\chi_i^*)", fontsize=15)

    add_paragraph_with_font(doc, "手順2: 予測観測値", size=10, bold=True)
    add_math_image(
        doc,
        r"\hat{\mathbf{z}} = \sum_{i=0}^{2n} W_i^m \mathbf{Z}_i",
        fontsize=15,
    )

    add_paragraph_with_font(doc, "手順3: イノベーション共分散", size=10, bold=True)
    add_math_image(
        doc,
        r"\mathbf{S} = \sum_{i=0}^{2n} W_i^c "
        r"(\mathbf{Z}_i - \hat{\mathbf{z}})(\mathbf{Z}_i - \hat{\mathbf{z}})^T + \mathbf{R}",
        fontsize=15,
    )

    add_paragraph_with_font(doc, "手順4: 交差共分散", size=10, bold=True)
    add_math_image(
        doc,
        r"\mathbf{P}_{xz} = \sum_{i=0}^{2n} W_i^c "
        r"(\chi_i^* - \hat{\mathbf{x}}^-)(\mathbf{Z}_i - \hat{\mathbf{z}})^T",
        fontsize=15,
    )

    add_paragraph_with_font(doc, "手順5: カルマンゲイン", size=10, bold=True)
    add_math_image(doc, r"\mathbf{K} = \mathbf{P}_{xz} \cdot \mathbf{S}^{-1}", fontsize=15)

    add_paragraph_with_font(doc, "手順6: 状態更新", size=10, bold=True)
    add_math_image(
        doc,
        r"\hat{\mathbf{x}} = \hat{\mathbf{x}}^- + \mathbf{K}(\mathbf{z} - \hat{\mathbf{z}})",
        fontsize=15,
    )

    add_paragraph_with_font(doc, "手順7: 共分散更新", size=10, bold=True)
    add_math_image(
        doc,
        r"\mathbf{P} = \mathbf{P}^- - \mathbf{K} \cdot \mathbf{S} \cdot \mathbf{K}^T",
        fontsize=15,
    )
    add_symbol_legend(doc, [
        ("Z_i", "観測空間に投影されたシグマポイント"),
        ("h(chi_i^*)", "観測モデル関数（距離・方位・仰角・ドップラーに変換）"),
        ("hat{z}", "予測観測値（シグマポイントの重み付き平均）"),
        ("S", "イノベーション共分散行列"),
        ("R", "観測ノイズ共分散行列"),
        ("P_{xz}", "状態-観測間クロス共分散行列"),
        ("K", "カルマンゲイン行列（= P_xz S^{-1}）"),
        ("z", "実際の観測値ベクトル"),
        ("hat{x}", "更新後の状態推定値"),
        ("P", "更新後の誤差共分散行列"),
        ("S^{-1}", "イノベーション共分散の逆行列"),
    ])

    # --- Process noise ---
    doc.add_heading("5.6 プロセスノイズ", level=2)
    add_paragraph_with_font(
        doc,
        "プロセスノイズ共分散行列 Q はデフォルトで以下の対角行列である（表5-2）。",
        size=10,
    )

    add_table_caption(doc, "表5-2  プロセスノイズパラメータ")
    add_table_with_header(
        doc,
        ["成分", "パラメータ", "デフォルト値"],
        [
            ("位置 (x, y, z)", "σ_pos", "160 m"),
            ("速度 (vx, vy, vz)", "σ_vel", "65 m/s"),
            ("加速度 (ax, ay, az)", "σ_acc", "160 m/s²"),
        ],
        col_widths=[5.0, 4.0, 4.0],
    )

    doc.add_paragraph()
    add_math_image(
        doc,
        r"\mathbf{Q} = \mathrm{diag}("
        r"\sigma_{pos}^2, \sigma_{pos}^2, \sigma_{pos}^2, "
        r"\sigma_{vel}^2, \sigma_{vel}^2, \sigma_{vel}^2, "
        r"\sigma_{acc}^2, \sigma_{acc}^2, \sigma_{acc}^2)",
        fontsize=13,
    )
    add_symbol_legend(doc, [
        ("Q", "プロセスノイズ共分散行列（9×9対角行列）"),
        ("sigma_{pos}", "位置成分のプロセスノイズ標準偏差（= 160 m）"),
        ("sigma_{vel}", "速度成分のプロセスノイズ標準偏差（= 65 m/s）"),
        ("sigma_{acc}", "加速度成分のプロセスノイズ標準偏差（= 160 m/s²）"),
        ("diag(...)", "対角行列コンストラクタ（非対角要素 = 0）"),
    ])

    doc.add_page_break()


def add_section_4(doc):
    """Section 4: IMM Filter."""
    doc.add_heading("6. IMM (Interacting Multiple Model) フィルタ", level=1)

    add_paragraph_with_font(
        doc,
        "IMM（Interacting Multiple Model）フィルタは、"
        "複数の運動モデルを同時に適用し、モデル確率に基づいて重み付け統合を行う。"
        "これにより、機動変化を伴う目標の追尾精度を向上させる。",
        size=10,
    )

    # --- Models ---
    doc.add_heading("6.1 運動モデル構成", level=2)
    add_paragraph_with_font(
        doc,
        "本システムでは3種類の運動モデルを使用する（表6-1）。",
        size=10,
    )

    add_table_caption(doc, "表6-1  IMMフィルタの運動モデル構成")
    add_table_with_header(
        doc,
        ["モデルID", "名称", "略称", "適用場面", "ノイズスケール"],
        [
            ("0", "等速直線", "CV", "巡航・中間飛翔段階", "10%（安定飛行）"),
            ("1", "弾道", "Ballistic", "弾道飛翔（重力＋大気抵抗）", "30%（物理モデルが正確）"),
            ("2", "旋回", "CT", "HGV機動・旋回飛翔", "100%（機動）"),
        ],
        col_widths=[2.0, 3.0, 2.5, 4.5, 4.0],
    )

    # --- Transition matrix ---
    doc.add_heading("6.2 モデル遷移確率行列", level=2)
    add_paragraph_with_font(
        doc,
        "モデル間の遷移はマルコフ連鎖でモデル化され、"
        "遷移確率行列 Π は表6-2のように定義される。",
        size=10,
    )

    add_table_caption(doc, "表6-2  モデル遷移確率行列 Π")
    add_table_with_header(
        doc,
        ["", "→ CV", "→ Ballistic", "→ CT"],
        [
            ("CV", "0.80", "0.15", "0.05"),
            ("Ballistic", "0.10", "0.85", "0.05"),
            ("CT", "0.05", "0.10", "0.85"),
        ],
        col_widths=[3.0, 3.5, 3.5, 3.5],
    )

    doc.add_paragraph()
    add_paragraph_with_font(
        doc,
        "初期モデル確率は一様分布: μ₀ = [1/3, 1/3, 1/3]",
        size=10,
    )

    # --- IMM Cycle ---
    doc.add_heading("6.3 IMMサイクル", level=2)

    add_paragraph_with_font(doc, "手順1: 混合（Mixing）", size=10, bold=True)
    add_paragraph_with_font(
        doc,
        "各モデルの予測確率を遷移確率と現在のモデル確率から計算する。",
        size=10,
    )
    add_math_image(
        doc,
        r"c_j = \sum_k \pi(k \to j) \cdot \mu_k",
        fontsize=15,
    )

    add_paragraph_with_font(doc, "手順2: モデル別予測", size=10, bold=True)
    add_paragraph_with_font(
        doc,
        "各運動モデルがそれぞれ独立に状態予測を行う（CV, Ballistic, CTモデル）。",
        size=10,
    )

    add_paragraph_with_font(doc, "手順3: 重み付け統合", size=10, bold=True)
    add_math_image(
        doc,
        r"\hat{\mathbf{x}} = \sum_j c_j \cdot \hat{\mathbf{x}}_j",
        fontsize=15,
    )
    add_math_image(
        doc,
        r"\mathbf{P} = \sum_j c_j \left[ \mathbf{P}_j + "
        r"(\hat{\mathbf{x}}_j - \hat{\mathbf{x}})(\hat{\mathbf{x}}_j - \hat{\mathbf{x}})^T \right]",
        fontsize=15,
    )

    add_paragraph_with_font(doc, "手順4: 尤度ベース更新（観測後）", size=10, bold=True)
    add_paragraph_with_font(
        doc,
        "観測値を取得した後、各モデルの尤度に基づきモデル確率を更新する。",
        size=10,
    )
    add_math_image(
        doc,
        r"L_j = \mathcal{N}(\mathrm{innovation}_j;\ 0,\ \mathbf{S}_j)",
        fontsize=15,
    )
    add_math_image(
        doc,
        r"\mu_j = \frac{c_j \cdot L_j}{\sum_k c_k \cdot L_k}",
        fontsize=15,
    )

    add_paragraph_with_font(
        doc,
        "数値安定性のため、尤度計算はlog空間で実行する。"
        "また、モデル確率の下限値（floor）を0.01に設定し、"
        "特定モデルの確率が0に収束することを防止する。",
        size=10,
    )
    add_symbol_legend(doc, [
        ("mu_k", "モデル k の現在の確率（更新前）"),
        ("pi(k->j)", "モデル k からモデル j への遷移確率"),
        ("c_j", "混合ステップでのモデル j の予測確率"),
        ("hat{x}", "IMM統合後の状態推定値"),
        ("hat{x}_j", "モデル j のUKF状態推定値"),
        ("P", "IMM統合後の誤差共分散行列"),
        ("P_j", "モデル j の誤差共分散行列"),
        ("L_j", "モデル j の尤度（イノベーションの正規分布確率密度）"),
        ("N(v; 0, S)", "平均0・共分散 S のガウス分布"),
        ("S_j", "モデル j のイノベーション共分散行列"),
        ("mu_j", "モデル j の更新後確率（規格化済み）"),
    ])

    doc.add_page_break()


def add_section_5(doc):
    """Section 5: Motion Models."""
    doc.add_heading("7. 運動モデル", level=1)

    # --- 5.1 CV ---
    doc.add_heading("7.1 CV（等速直線）モデル", level=2)
    add_paragraph_with_font(
        doc,
        "巡航・中間飛翔段階の安定した飛行を表現するモデルである。"
        "加速度は時定数 τ=5s で指数減衰する。",
        size=10,
    )

    add_paragraph_with_font(doc, "位置更新:", size=10, bold=True)
    add_math_image(doc, r"\mathbf{p}(t+\Delta t) = \mathbf{p}(t) + \mathbf{v} \cdot \Delta t", fontsize=15)

    add_paragraph_with_font(doc, "速度更新:", size=10, bold=True)
    add_math_image(doc, r"\mathbf{v}(t+\Delta t) = \mathbf{v}(t)", fontsize=15)

    add_paragraph_with_font(doc, "加速度更新（指数減衰）:", size=10, bold=True)
    add_math_image(
        doc,
        r"\mathbf{a}(t+\Delta t) = \mathbf{a}(t) \cdot \exp\left(-\frac{\Delta t}{\tau}\right)"
        r"\quad (\tau = 5\ \mathrm{s})",
        fontsize=15,
    )
    add_symbol_legend(doc, [
        ("p(t)", "時刻 t における位置ベクトル [m]"),
        ("v(t)", "時刻 t における速度ベクトル [m/s]"),
        ("a(t)", "時刻 t における加速度ベクトル [m/s²]"),
        ("Delta t", "タイムステップ [s]"),
        ("tau", "加速度の指数減衰時定数（= 5 s）"),
        ("exp(...)", "指数関数"),
    ])

    # --- 5.2 Ballistic ---
    doc.add_heading("7.2 弾道（RK4）モデル", level=2)
    add_paragraph_with_font(
        doc,
        "重力および大気抵抗を考慮した物理ベースの弾道運動モデルである。"
        "RK4（4次ルンゲ・クッタ法）による数値積分を用いる。",
        size=10,
    )

    add_paragraph_with_font(doc, "高度依存重力:", size=10, bold=True)
    add_math_image(
        doc,
        r"g(h) = g_0 \cdot \left(\frac{R_E}{R_E + h}\right)^2",
        fontsize=15,
    )

    add_table_caption(doc, "表7-1  重力定数")
    add_table_with_header(
        doc,
        ["定数", "記号", "値"],
        [
            ("標準重力加速度", "g₀", "9.80665 m/s²"),
            ("地球半径", "R_E", "6,371,000 m"),
        ],
        col_widths=[5.0, 3.0, 5.0],
    )

    doc.add_paragraph()
    add_paragraph_with_font(doc, "大気密度モデル（指数大気）:", size=10, bold=True)
    add_math_image(
        doc,
        r"\rho(h) = \rho_0 \cdot \exp\left(-\frac{h}{H}\right)",
        fontsize=15,
    )

    add_table_caption(doc, "表7-2  大気密度モデルパラメータ")
    add_table_with_header(
        doc,
        ["定数", "記号", "値"],
        [
            ("海面大気密度", "ρ₀", "1.225 kg/m³"),
            ("スケールハイト", "H", "7,400 m"),
        ],
        col_widths=[5.0, 3.0, 5.0],
    )

    doc.add_paragraph()
    add_paragraph_with_font(doc, "抗力:", size=10, bold=True)
    add_math_image(
        doc,
        r"F_{drag} = \beta \cdot \rho(h) \cdot |\mathbf{v}| \cdot \mathbf{v}",
        fontsize=15,
    )
    add_paragraph_with_font(
        doc,
        "ここで β = Cd·A/(2m) = 0.001 は代表的な弾道係数である。",
        size=10,
    )

    add_paragraph_with_font(doc, "運動方程式:", size=10, bold=True)
    add_math_image(doc, r"\ddot{x} = -\beta \cdot \rho \cdot |\mathbf{v}| \cdot v_x", fontsize=15)
    add_math_image(doc, r"\ddot{y} = -\beta \cdot \rho \cdot |\mathbf{v}| \cdot v_y", fontsize=15)
    add_math_image(
        doc,
        r"\ddot{z} = -g(h) - \beta \cdot \rho \cdot |\mathbf{v}| \cdot v_z",
        fontsize=15,
    )

    add_paragraph_with_font(
        doc,
        "上記の微分方程式をRK4（4段ルンゲ・クッタ法）で各タイムステップごとに数値積分する。",
        size=10,
    )
    add_symbol_legend(doc, [
        ("g(h)", "高度 h における重力加速度 [m/s²]"),
        ("g_0", "標準重力加速度 = 9.80665 m/s²"),
        ("R_E", "地球半径 = 6,371,000 m"),
        ("h", "目標の高度（= z） [m]"),
        ("rho(h)", "高度 h における大気密度 [kg/m³]"),
        ("rho_0", "海面大気密度 = 1.225 kg/m³"),
        ("H", "大気スケール高度 = 7,400 m"),
        ("F_{drag}", "抗力ベクトル [N]"),
        ("beta", "弾道係数（= C_d A / (2m) = 0.001 m²/kg）"),
        ("|v|", "速度の大きさ [m/s]"),
        ("ddot{x}, ddot{y}", "水平面内の加速度成分（抗力のみ） [m/s²]"),
        ("ddot{z}", "鉛直方向加速度（重力＋抗力） [m/s²]"),
    ])

    # --- 5.3 CT ---
    doc.add_heading("7.3 CT（旋回）モデル", level=2)
    add_paragraph_with_font(
        doc,
        "HGVの機動飛翔（旋回）を表現するCoordinated Turnモデルである。"
        "旋回角速度 ω を状態から推定し、旋回運動を適用する。",
        size=10,
    )

    add_paragraph_with_font(doc, "旋回角速度の推定:", size=10, bold=True)
    add_math_image(
        doc,
        r"\omega = \frac{v_x \cdot a_y - v_y \cdot a_x}{v_x^2 + v_y^2}",
        fontsize=15,
    )
    add_paragraph_with_font(
        doc,
        "ω は ±0.785 rad/s（±45°/s）にクランプされる。",
        size=10,
    )

    add_paragraph_with_font(doc, "|ω| > 10⁻⁴ の場合（旋回）:", size=10, bold=True)
    add_paragraph_with_font(doc, "水平面内の位置更新:", size=10)
    add_math_image(
        doc,
        r"x(t+\Delta t) = x + \frac{v_x \sin(\omega \Delta t) + v_y (\cos(\omega \Delta t) - 1)}{\omega}",
        fontsize=14,
    )
    add_math_image(
        doc,
        r"y(t+\Delta t) = y + \frac{-v_x (\cos(\omega \Delta t) - 1) + v_y \sin(\omega \Delta t)}{\omega}",
        fontsize=14,
    )

    add_paragraph_with_font(doc, "水平面内の速度更新:", size=10)
    add_math_image(
        doc,
        r"v_x' = v_x \cos(\omega \Delta t) - v_y \sin(\omega \Delta t)",
        fontsize=14,
    )
    add_math_image(
        doc,
        r"v_y' = v_x \sin(\omega \Delta t) + v_y \cos(\omega \Delta t)",
        fontsize=14,
    )

    add_paragraph_with_font(doc, "向心加速度:", size=10)
    add_math_image(
        doc,
        r"a_x = -\omega \cdot v_y', \quad a_y = \omega \cdot v_x'",
        fontsize=14,
    )

    add_paragraph_with_font(doc, "|ω| ≤ 10⁻⁴ の場合（直線近似）:", size=10, bold=True)
    add_paragraph_with_font(
        doc,
        "旋回角速度が十分小さい場合、等加速度直線運動モデルにフォールバックする。",
        size=10,
    )

    add_paragraph_with_font(doc, "垂直方向（共通）:", size=10, bold=True)
    add_math_image(
        doc,
        r"z(t+\Delta t) = z + v_z \cdot \Delta t + \frac{1}{2} a_z \cdot \Delta t^2",
        fontsize=15,
    )
    add_symbol_legend(doc, [
        ("omega", "旋回角速度（= (v_x a_y - v_y a_x) / (v_x²+v_y²)） [rad/s]"),
        ("v_x, v_y", "水平面内の速度成分 [m/s]"),
        ("a_x, a_y", "水平面内の加速度成分 [m/s²]"),
        ("v_x', v_y'", "旋回後の速度成分 [m/s]"),
        ("x, y", "水平位置成分 [m]"),
        ("Delta t", "タイムステップ [s]"),
        ("z", "高度 [m]"),
        ("v_z", "鉛直速度成分 [m/s]"),
        ("a_z", "鉛直加速度成分 [m/s²]"),
        ("sin, cos", "三角関数"),
    ])

    doc.add_page_break()


def add_section_6(doc):
    """Section 6: Data Association."""
    doc.add_heading("8. データアソシエーション", level=1)

    add_paragraph_with_font(
        doc,
        "データアソシエーションは、レーダ観測値を既存航跡に対応付ける処理である。"
        "コスト行列の計算はGPU上で並列実行され、最適割当はハンガリアン法（Munkres法）で解く。",
        size=10,
    )

    # --- Cost matrix ---
    doc.add_heading("8.1 コスト行列", level=2)
    add_paragraph_with_font(
        doc,
        "コスト行列の各要素は、正規化イノベーション距離（NID）で計算される。"
        "GPU上で各観測-航跡ペアを並列に計算する。",
        size=10,
    )
    add_math_image(
        doc,
        r"d^2(i,j) = \sum_{k \in \{r, az, el, dop\}} "
        r"\left(\frac{z_k^{meas} - z_k^{pred}}{\sigma_k}\right)^2",
        fontsize=15,
    )
    add_symbol_legend(doc, [
        ("d^2(i,j)", "航跡 i と観測 j の正規化イノベーション距離（NID）"),
        ("k", "観測成分インデックス（r: 距離, az: 方位, el: 仰角, dop: ドップラー）"),
        ("z_k^{meas}", "観測値の第 k 成分"),
        ("z_k^{pred}", "航跡 i の予測観測値の第 k 成分"),
        ("sigma_k", "第 k 観測成分のノイズ標準偏差"),
        ("i", "航跡インデックス"),
        ("j", "観測インデックス"),
    ])

    # --- Gating ---
    doc.add_heading("8.2 ゲーティング", level=2)
    add_paragraph_with_font(
        doc,
        "計算された距離がゲート閾値を超える場合、当該ペアを棄却する（表8-1）。",
        size=10,
    )

    add_table_caption(doc, "表8-1  ゲーティングパラメータ")
    add_table_with_header(
        doc,
        ["パラメータ", "値", "説明"],
        [
            ("ゲート閾値", "500.0", "NIDがこの値を超えるペアを棄却"),
            ("棄却コスト", "10¹⁰", "棄却されたペアに割り当てるコスト値"),
        ],
        col_widths=[4.0, 3.0, 9.0],
    )

    # --- Hungarian ---
    doc.add_heading("8.3 ハンガリアン法（Munkres法）", level=2)
    add_paragraph_with_font(
        doc,
        "ハンガリアン法は O(n³) の計算量で最適割当問題を解くアルゴリズムである。"
        "以下のステップで構成される。",
        size=10,
    )

    steps = [
        ("行縮約", "各行の最小値を行全体から減算する"),
        ("列縮約", "各列の最小値を列全体から減算する"),
        ("初期スターリング", "各行・列で唯一のゼロ要素にスターを付与する"),
        ("列カバー", "スター付きゼロを含む列をカバーする"),
        ("未カバーゼロの探索", "未カバーのゼロ要素にプライムを付与する"),
        ("拡張パス", "スターとプライムの交互パスにより割当を改善する"),
        ("行列調整", "未カバー要素の最小値で行列を調整し、手順4に戻る"),
    ]
    for i, (step_name, step_desc) in enumerate(steps, 1):
        add_paragraph_with_font(
            doc,
            f"ステップ{i}: {step_name} — {step_desc}",
            size=10,
        )

    doc.add_page_break()


def add_section_7(doc):
    """Section 7: Track Management."""
    doc.add_heading("9. 航跡管理", level=1)

    add_paragraph_with_font(
        doc,
        "航跡管理は、航跡のライフサイクル（生成・確認・消失）を制御する。",
        size=10,
    )

    # --- States ---
    doc.add_heading("9.1 航跡状態", level=2)
    add_paragraph_with_font(
        doc,
        "航跡は表9-1に示す3状態で管理される。",
        size=10,
    )

    add_table_caption(doc, "表9-1  航跡状態の定義")
    add_table_with_header(
        doc,
        ["状態", "英語名", "説明", "遷移条件"],
        [
            ("仮航跡", "TENTATIVE", "新規生成された未確認航跡", "未対応観測から生成"),
            ("確認航跡", "CONFIRMED", "信頼性が確認された航跡", "ヒット数 ≥ confirm_hits (=2)"),
            ("消失航跡", "LOST", "観測が途絶えた航跡", "ミス数 ≥ delete_misses (=90 ≈ 3s @30Hz)"),
        ],
        col_widths=[2.5, 3.0, 5.0, 5.5],
    )

    doc.add_paragraph()
    add_paragraph_with_font(
        doc,
        "状態遷移: TENTATIVE → CONFIRMED → LOST",
        size=10,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )

    # --- Initialization ---
    doc.add_heading("9.2 航跡初期化", level=2)
    add_paragraph_with_font(
        doc,
        "未対応の観測値から新規航跡を初期化する。極座標観測を直交座標に変換する。",
        size=10,
    )

    add_paragraph_with_font(doc, "位置の初期化（極座標→直交座標）:", size=10, bold=True)
    add_math_image(
        doc,
        r"x = r \cos(\phi) \cos(\theta) + x_{sensor}",
        fontsize=14,
    )
    add_math_image(
        doc,
        r"y = r \cos(\phi) \sin(\theta) + y_{sensor}",
        fontsize=14,
    )
    add_math_image(
        doc,
        r"z = r \sin(\phi) + z_{sensor}",
        fontsize=14,
    )
    add_symbol_legend(doc, [
        ("x, y, z", "直交座標（東・北・高度） [m]"),
        ("r", "スラントレンジ（観測距離） [m]"),
        ("phi", "仰角 [rad]"),
        ("theta", "方位角 [rad]"),
        ("x_{sensor}, y_{sensor}, z_{sensor}", "センサの位置座標 [m]"),
        ("sin, cos", "三角関数"),
    ])

    add_paragraph_with_font(
        doc,
        "速度はドップラー観測値に基づき視線方向（LOS）に沿って初期化される。"
        "初期共分散は観測距離に応じてスケーリングされる。",
        size=10,
    )

    # --- Parameters ---
    doc.add_heading("9.3 航跡管理パラメータ", level=2)
    add_paragraph_with_font(doc, "航跡管理に使用するパラメータを表9-2に示す。", size=10)
    add_table_caption(doc, "表9-2  航跡管理パラメータ")
    add_table_with_header(
        doc,
        ["パラメータ", "デフォルト値", "説明"],
        [
            ("confirm_hits", "2", "確認に必要なヒット数"),
            ("delete_misses", "90", "削除までのミス数（30Hzで約3秒）"),
            ("初期化最小SNR", "22.0 dB", "航跡初期化に必要な最小SNR"),
        ],
        col_widths=[4.0, 3.5, 8.5],
    )

    doc.add_page_break()


def add_section_8(doc):
    """Section 8: Evaluation Metrics."""
    doc.add_heading("10. 評価指標", level=1)

    add_paragraph_with_font(
        doc,
        "追尾性能を定量的に評価するための各種指標を以下に定義する。",
        size=10,
    )

    # --- RMSE ---
    doc.add_heading("10.1 RMSE（二乗平均平方根誤差）", level=2)

    add_paragraph_with_font(doc, "位置RMSE:", size=10, bold=True)
    add_math_image(
        doc,
        r"\mathrm{RMSE}_{pos} = \sqrt{\frac{1}{N} \sum_{i=1}^{N} "
        r"\|\mathbf{p}_{track,i} - \mathbf{p}_{truth,i}\|^2}",
        fontsize=15,
    )

    add_paragraph_with_font(doc, "速度RMSE:", size=10, bold=True)
    add_math_image(
        doc,
        r"\mathrm{RMSE}_{vel} = \sqrt{\frac{1}{N} \sum_{i=1}^{N} "
        r"\|\mathbf{v}_{track,i} - \mathbf{v}_{truth,i}\|^2}",
        fontsize=15,
    )
    add_symbol_legend(doc, [
        ("RMSE_{pos}", "位置の二乗平均平方根誤差 [m]"),
        ("RMSE_{vel}", "速度の二乗平均平方根誤差 [m/s]"),
        ("N", "評価サンプル数（タイムステップ数）"),
        ("p_{track,i}", "タイムステップ i の追尾位置ベクトル [m]"),
        ("p_{truth,i}", "タイムステップ i の真値位置ベクトル [m]"),
        ("v_{track,i}", "タイムステップ i の追尾速度ベクトル [m/s]"),
        ("v_{truth,i}", "タイムステップ i の真値速度ベクトル [m/s]"),
        ("||...||", "ユークリッドノルム（3次元距離）"),
        ("sqrt{...}", "平方根"),
        ("sum_{i=1}^N", "i = 1 から N までの総和"),
    ])

    # --- OSPA ---
    doc.add_heading("10.2 OSPA（Optimal SubPattern Assignment）", level=2)
    add_paragraph_with_font(
        doc,
        "OSPAはパターン間の距離指標であり、位置誤差と個数誤差を統合的に評価する。"
        "OSPAパラメータを表10-1に示す。",
        size=10,
    )
    add_math_image(
        doc,
        r"d_{OSPA}^{(p)}(X,Y) = \left[ \frac{1}{\max(m,n)} "
        r"\left( \sum_{i} \min(d(x_i, y_{\pi(i)}), c)^p + |m-n| \cdot c^p \right) "
        r"\right]^{1/p}",
        fontsize=14,
    )
    add_symbol_legend(doc, [
        ("d_{OSPA}^{(p)}(X,Y)", "集合 X（追尾）と Y（真値）間のOSPA距離"),
        ("p", "ノルム次数（= 1）"),
        ("c", "カットオフ距離（= 10,000 m）。個数誤差のペナルティ上限"),
        ("m", "追尾集合 X の要素数（航跡数）"),
        ("n", "真値集合 Y の要素数（目標数）"),
        ("x_i, y_j", "集合 X, Y の各要素（位置ベクトル）"),
        ("pi(i)", "最適割当によるインデックスの対応付け"),
        ("d(x_i, y_j)", "x_i と y_j のユークリッド距離 [m]"),
        ("min, max", "最小値・最大値演算"),
        ("|m-n|", "追尾数と目標数の差（個数誤差）"),
    ])

    add_table_caption(doc, "表10-1  OSPAパラメータ")
    add_table_with_header(
        doc,
        ["パラメータ", "記号", "デフォルト値", "説明"],
        [
            ("カットオフ距離", "c", "10,000 m", "個数ペナルティの上限距離"),
            ("次数", "p", "1", "距離のLpノルム次数"),
        ],
        col_widths=[3.5, 2.0, 3.5, 7.0],
    )

    # --- Detection metrics ---
    doc.add_heading("10.3 検出指標", level=2)
    add_paragraph_with_font(
        doc,
        "追尾の正確性を評価する検出指標を以下に定義する。各指標の意味を表10-2に示す。",
        size=10,
    )

    add_math_image(
        doc,
        r"\mathrm{Precision} = \frac{TP}{TP + FP}",
        fontsize=15,
    )
    add_math_image(
        doc,
        r"\mathrm{Recall} = \frac{TP}{TP + FN}",
        fontsize=15,
    )
    add_math_image(
        doc,
        r"F_1 = \frac{2 \cdot \mathrm{Precision} \cdot \mathrm{Recall}}"
        r"{\mathrm{Precision} + \mathrm{Recall}}",
        fontsize=15,
    )
    add_symbol_legend(doc, [
        ("TP (True Positive)", "正しく追尾された目標数（航跡と真値が対応）"),
        ("FP (False Positive)", "誤追尾数（真値に対応しない余剰な航跡）"),
        ("FN (False Negative)", "追尾漏れ数（航跡が対応しない未追尾の真値）"),
        ("Precision", "適合率（追尾航跡のうち正しいものの割合）"),
        ("Recall", "再現率（真値目標のうち追尾されたものの割合）"),
        ("F_1", "F1スコア（適合率と再現率の調和平均）"),
    ])

    add_table_caption(doc, "表10-2  検出指標の定義")
    add_table_with_header(
        doc,
        ["指標", "定義", "説明"],
        [
            ("TP (True Positive)", "正しく追尾された目標", "航跡と真値が対応"),
            ("FP (False Positive)", "誤追尾", "真値に対応しない航跡"),
            ("FN (False Negative)", "追尾漏れ", "航跡が対応しない真値"),
        ],
        col_widths=[4.0, 4.5, 7.5],
    )

    # --- Track continuity ---
    doc.add_heading("10.4 航跡連続性指標", level=2)
    add_paragraph_with_font(
        doc,
        "各真値目標に対する航跡の連続性を表10-3に示す3カテゴリで評価する。",
        size=10,
    )

    add_table_caption(doc, "表10-3  航跡連続性指標")
    add_table_with_header(
        doc,
        ["カテゴリ", "英語名", "条件", "説明"],
        [
            ("主追尾", "Mostly Tracked (MT)", "追尾率 ≥ 80%", "ライフタイムの80%以上で追尾"),
            ("部分追尾", "Partially Tracked (PT)", "20% ≤ 追尾率 < 80%", "ライフタイムの20～80%で追尾"),
            ("主消失", "Mostly Lost (ML)", "追尾率 < 20%", "ライフタイムの20%未満で追尾"),
        ],
        col_widths=[2.5, 4.0, 3.5, 6.0],
    )

    doc.add_page_break()


def add_section_gpu_acceleration(doc):
    """Section 11: GPU Parallel Processing."""
    doc.add_heading("11. GPU並列処理による高速化", level=1)

    add_paragraph_with_font(
        doc,
        "FastTrackerはNVIDIA CUDAを活用し、UKFシグマポイント演算・IMM並列処理・"
        "コスト行列計算を GPU 上で並列実行する。"
        "これにより多数の目標を同時にリアルタイム追尾することが可能となる。"
        "GPU並列化アーキテクチャの概要を図11-1に示す。",
        size=10,
    )

    fig_gpu = _make_gpu_arch_figure()
    fig_gpu_path = save_figure(fig_gpu, "gpu_arch.png")
    add_figure_to_doc(
        doc, fig_gpu_path,
        caption="図11-1  GPU並列処理アーキテクチャ",
        width_inches=6.5,
    )

    # 11.1 UKF sigma point parallelism
    doc.add_heading("11.1 UKFシグマポイントの並列化", level=2)
    add_paragraph_with_font(
        doc,
        "UKFは状態次元 n=9 に対し 2n+1=19 個のシグマポイントを用いる。"
        "N 個の目標に対するシグマポイント生成・予測・観測モデル計算は、"
        "N × 19 個のスレッドが同時実行するよう設計されている。"
        "主要 CUDAカーネルと並列化戦略を表11-1に示す。",
        size=10,
    )

    add_table_caption(doc, "表11-1  UKF CUDAカーネル一覧")
    add_table_with_header(
        doc,
        ["カーネル名", "グリッド／ブロック構成", "処理内容"],
        [
            ("generateSigmaPoints",
             "grid = (N×19+255)/256\nblock = 256",
             "コレスキー分解済み P から全目標のシグマポイントを生成"),
            ("predictSigmaPoints",
             "grid = (N×19+255)/256\nblock = 256",
             "各シグマポイントに対して運動モデル（CV/弾道/CT）を適用"),
            ("measurementModel",
             "grid = (N×19+255)/256\nblock = 256",
             "シグマポイントを観測空間（距離・方位角・仰角・ドップラー）に投影"),
            ("computeWeightedMean",
             "grid = N\nblock = 19",
             "重み付き平均 x̄, z̄ を共有メモリでリダクション計算"),
            ("computeCovariance",
             "grid = N\nblock = 19",
             "共分散 P_pred および S（観測共分散）を計算"),
            ("computeCrossCovariance",
             "grid = N\nblock = 19",
             "クロス共分散行列 P_xz を計算"),
            ("computeKalmanGain",
             "grid = N\nblock = 1",
             "カルマンゲイン K = P_xz · S⁻¹ を計算（S の逆行列は CPU）"),
            ("updateState",
             "grid = N\nblock = 1",
             "状態更新 x̂ = x̄ + K·(z − z̄) を適用"),
            ("updateCovariance",
             "grid = N\nblock = 1",
             "共分散更新 P = P_pred − K·S·Kᵀ を適用"),
        ],
        col_widths=[4.0, 4.5, 7.5],
    )

    add_paragraph_with_font(
        doc,
        "物理定数（g₀, R_E, ρ₀, H, BALLISTIC_BETA 等）はすべて "
        "__constant__ メモリに配置され、全スレッドからキャッシュを通じて高速にアクセスできる。",
        size=10,
    )
    add_math_image(doc,
        r"\text{Throughput} \propto \frac{N_{targets} \times 19}{\text{blockDim}=256}",
        fontsize=13)
    add_symbol_legend(doc, [
        ("N_{targets}", "追尾対象の目標数（並列処理のバッチサイズ）"),
        ("19", "シグマポイント総数（= 2n+1 = 2×9+1）"),
        ("blockDim", "CUDAスレッドブロックサイズ（= 256スレッド）"),
        ("Throughput", "GPU処理スループット（ブロック数 ∝ N × 19 / 256）"),
        ("N_targets x 19", "1ステップ全シグマポイント総数"),
    ])

    # 11.2 IMM parallel streams
    doc.add_heading("11.2 IMMフィルタのストリーム並列化", level=2)
    add_paragraph_with_font(
        doc,
        "IMMフィルタは CV・弾道・CT の3モデルを管理する。"
        "各モデルのUKF predict/updateは独立した CUDAストリームで同時実行される。"
        "これにより3モデルの逐次実行に比べ、最大3倍の並列度を達成する。"
        "IMM並列化の処理フローを表11-2に示す。",
        size=10,
    )

    add_table_caption(doc, "表11-2  IMM CUDAストリーム並列化フロー")
    add_table_with_header(
        doc,
        ["ステップ", "処理", "実行ストリーム"],
        [
            ("1. 混合初期化",
             "computeMixingProbabilitiesKernel\n混合確率 μ_ij と混合推定値 x̃₀ⱼ を計算",
             "Stream 0（同期）"),
            ("2. モデル並列 UKF",
             "各モデル（CV, 弾道, CT）の UKF predict + update を実行",
             "Stream 0 / 1 / 2\n（3ストリーム並列）"),
            ("3. cudaDeviceSynchronize",
             "全ストリームの完了を待機",
             "—"),
            ("4. 確率更新",
             "updateModelProbabilitiesKernel\n尤度 Λⱼ からモデル確率 μⱼ を更新",
             "Stream 0（同期）"),
            ("5. 推定統合",
             "combineEstimatesKernel\n最終推定値 x̂ = Σⱼ μⱼ x̂ⱼ を計算",
             "Stream 0（同期）"),
        ],
        col_widths=[3.0, 7.5, 5.5],
    )

    # 11.3 Cost matrix parallelism
    doc.add_heading("11.3 コスト行列計算の並列化", level=2)
    add_paragraph_with_font(
        doc,
        "データアソシエーションでは N_tracks × N_meas 個のマハラノビス距離を計算する必要がある。"
        "各要素 (i, j) を独立したスレッドで並列計算することで、"
        "逐次ループと比べてほぼ線形なスループット向上を実現する。",
        size=10,
    )
    add_math_image(doc,
        r"d^2_{ij} = (\mathbf{z}_j - \hat{\mathbf{z}}_i)^\top S_i^{-1} (\mathbf{z}_j - \hat{\mathbf{z}}_i)",
        fontsize=14)
    add_symbol_legend(doc, [
        ("d^2_{ij}", "航跡 i と観測 j 間のマハラノビス距離の二乗"),
        ("z_j", "j 番目の観測値ベクトル（距離・方位・仰角・ドップラー）"),
        ("hat{z}_i", "航跡 i の予測観測値ベクトル"),
        ("S_i", "航跡 i のイノベーション共分散行列（観測空間4×4）"),
        ("S_i^{-1}", "S_i の逆行列（ホワイトニング変換）"),
        ("T (上付き)", "ベクトル転置"),
        ("i", "航跡インデックス（i = 1 ... N_tracks）"),
        ("j", "観測インデックス（j = 1 ... N_meas）"),
    ])
    add_paragraph_with_font(
        doc,
        "ここで z_j は j 番目の観測値、ẑ_i は i 番目の航跡の予測観測値、"
        "S_i は革新共分散行列。ゲーティング閾値 d² > 500.0 の組み合わせは"
        "コスト行列から除外し，ハンガリアン法（Munkres）の処理量を削減する。",
        size=10,
    )

    add_table_caption(doc, "表11-3  データアソシエーション GPU設定")
    add_table_with_header(
        doc,
        ["設定項目", "値", "説明"],
        [
            ("コスト行列カーネルスレッド数", "blockDim = 256",
             "1ブロックあたりのスレッド数"),
            ("グリッドサイズ", "(N_tracks×N_meas+255)/256",
             "全要素を網羅するブロック数"),
            ("ゲーティング閾値", "d² = 500.0",
             "マハラノビス距離の二乗によるゲーティング"),
            ("__constant__メモリ対象", "g₀, R_E, ρ₀, H, BALLISTIC_BETA",
             "全カーネルから参照される物理定数"),
        ],
        col_widths=[5.0, 3.5, 7.5],
    )

    doc.add_page_break()


def add_revision_history(doc):
    """Add revision history table."""
    doc.add_heading("変更履歴", level=1)

    add_table_with_header(
        doc,
        ["版数", "日付", "変更内容", "担当"],
        [
            ("1.0", "2026/02/17", "初版作成", "—"),
            ("", "", "", ""),
            ("", "", "", ""),
        ],
        col_widths=[2.0, 3.0, 8.0, 3.0],
    )


# ---------------------------------------------------------------------------
# Main generation
# ---------------------------------------------------------------------------

def generate_document():
    """Generate the complete algorithm description document."""
    print("FastTracker Algorithm Document Generator")
    print("=" * 50)

    doc = Document()

    print("[1/14] Setting up page layout and styles...")
    setup_page(doc)
    setup_styles(doc)

    print("[2/14] Creating title page...")
    add_title_page(doc)

    print("[3/14] Adding table of contents...")
    add_toc(doc)

    print("[4/14] Section 1: System Overview...")
    add_section_1(doc)

    print("[5/14] Section 2: Target Generator...")
    add_section_target_generator(doc)

    print("[6/14] Section 3: Radar Simulator...")
    add_section_radar_simulator(doc)

    print("[7/14] Section 4: State-Space Model...")
    add_section_2(doc)

    print("[8/14] Section 5: UKF...")
    add_section_3(doc)

    print("[9/14] Section 6: IMM Filter...")
    add_section_4(doc)

    print("[10/14] Section 7: Motion Models...")
    add_section_5(doc)

    print("[11/14] Section 8-9: Data Association & Track Management...")
    add_section_6(doc)
    add_section_7(doc)

    print("[12/14] Section 10: Evaluation Metrics...")
    add_section_8(doc)

    print("[13/14] Section 11: GPU Acceleration...")
    add_section_gpu_acceleration(doc)

    print("[14/14] Adding revision history...")
    add_revision_history(doc)

    # Save
    print(f"\nSaving document to: {OUTPUT_DOCX}")
    doc.save(OUTPUT_DOCX)
    print("Document saved successfully.")

    # Cleanup temp images
    try:
        shutil.rmtree(TEMP_DIR)
        print(f"Cleaned up temporary files: {TEMP_DIR}")
    except Exception as e:
        print(f"Warning: Could not clean up temp dir: {e}")

    print(f"\nGenerated: {OUTPUT_DOCX}")
    print(f"Math images rendered: {_image_counter}")
    return OUTPUT_DOCX


if __name__ == "__main__":
    generate_document()
